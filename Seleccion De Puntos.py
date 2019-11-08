import pyodbc
import subprocess
from docx import Document
from sys import argv
from docx.shared import Pt
from re import sub,search
from time import mktime,strptime,strftime,localtime
from rutas import plantillaSeleccion,db
from wordExe import wordExe
from Informacion import manejarErrores,choices,cantidad

#VARIABLES
diasParaAtras = 7
plantilla = Document(plantillaSeleccion)
if len(argv)>1 and argv[1] == 'datecap':
	print('DATECAP')
	query = 'SELECT DISTINCT [Seleccion de Puntos].[Pto de medición], [Seleccion de Puntos].[Tipo Punto], [Seleccion de Puntos].[Fecha Instalacion], suministros_sidac.ESTADO, [Seleccion de Puntos].Reclamo, [Seleccion de Puntos].[Nro de Orden_ESJ], Usuarios.TIPO, Usuarios.TARIFA, suministros_sidac.SETA, suministros_sidac.LOCALIDAD FROM ([Seleccion de Puntos] LEFT JOIN suministros_sidac ON [Seleccion de Puntos].[Pto de medición] = suministros_sidac.CLIENTE) LEFT JOIN Usuarios ON [Seleccion de Puntos].[Pto de medición] = Usuarios.NUMERO_DE_SUMINISTRO WHERE [Seleccion de Puntos].[Fecha Instalacion] > DATEVALUE(?) and [Seleccion de Puntos].[Fecha Instalacion] < DATE();'
else:
	query = 'SELECT DISTINCT [Seleccion de Puntos].[Pto de medición], [Seleccion de Puntos].[Tipo Punto], [Seleccion de Puntos].[Fecha Instalacion], suministros_sidac.ESTADO, [Seleccion de Puntos].Reclamo, [Seleccion de Puntos].[Nro de Orden_ESJ], Usuarios.TIPO, Usuarios.TARIFA, suministros_sidac.SETA, suministros_sidac.LOCALIDAD FROM ([Seleccion de Puntos] LEFT JOIN suministros_sidac ON [Seleccion de Puntos].[Pto de medición] = suministros_sidac.CLIENTE) LEFT JOIN Usuarios ON [Seleccion de Puntos].[Pto de medición] = Usuarios.NUMERO_DE_SUMINISTRO WHERE [Seleccion de Puntos].[Fecha Instalacion] > DATEVALUE(?);'

def check(listaDePuntos):
	seleccionDeHoy = []
	seleccionDeAntes = []
	REPETIDO = False
	REMEDICION_TEMPRANA = False
	
	for punto in listaDePuntos:
		if punto[2].strftime('%d/%m/%Y') == fecha:
			seleccionDeHoy.append(punto)
		else: seleccionDeAntes.append(punto)
	try:	
		seleccionDeHoy = sorted(seleccionDeHoy,key=lambda x: x[5])
	except:
		print("Te olvidaste de copiar la columna de orden")
		input()
		exit()
	suministrosDeHoy = list(set([x[0] for x in seleccionDeHoy]))
	for sum in suministrosDeHoy:
		for med in seleccionDeAntes:
			if sum == med[0]:
				if fechaSeconds-(med[2]-med[2].utcfromtimestamp(0)).total_seconds() < 604800.0:
					print('El sum {} está siendo medido desde el {} | {} | {}'.format(sum,med[2].strftime('%d/%m/%Y'),med[1],int(med[5])))
					REPETIDO = True
				else:
					print('El sum {} se seleccionó el {} | {},{}'.format(sum,med[2].strftime('%d/%m/%Y'),med[1],med[5]))
					REMEDICION_TEMPRANA = True
	if REPETIDO:
		choice = input('Desea continuar de todas maneras? S/n :> ')
		if choice.lower() not in choices['yes']:
			exit()
	if REMEDICION_TEMPRANA:
		choice = input('¿Continuar? S/n :> ')
		if choice.lower() not in choices['yes']:
			exit()
		
	if len(seleccionDeHoy) > 32:
		print(len(seleccionDeHoy))
		print('\nHay más de 16 puntos. Posiblemente está mal la fecha de instalación o haya un suministro repetido en la tabla Usuarios')
		input()
		exit()
	puntos = [x[0] for x in seleccionDeHoy[::2]]
	if len(puntos) != len(set(puntos)):
		print('\nHay puntos repetidos para la selección de este día')
		input()
		exit()
	# seleccionDeHoy = sorted(seleccionDeHoy,key=lambda x: x[])
	return seleccionDeHoy

@manejarErrores
def main():
	error = False
	global fecha,fechaSeconds
	print('Fecha de instalacion:> xx/xx/xxxx\r',end='Fecha de instalacion:> ')
	fecha = input()
	if not search('^\d\d/\d\d/\d\d\d\d$',fecha):
		print('\nFormato de fecha erroneo. Este programita es medio tonto y le tenes que poner la fecha asi:\ndd/mm/yyyy')
		input()
		exit()
	fechaTupla = strptime(fecha,'%d/%m/%Y')
	fechaSeconds = mktime(fechaTupla)
	fechaInicioAnalisisSeconds = fechaSeconds - diasParaAtras*24*60*60.0
	fechaInicioAnalisis = strftime('%d/%m/%Y',localtime(fechaInicioAnalisisSeconds))

	conn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %db)
	cursor = conn.cursor()
	cursor.execute(query,(fechaInicioAnalisis))
	lista = cursor.fetchall()
	lista = check(lista)

	#TABLA DE PUNTOS
	puntero = 1
	tablaDePuntos = plantilla.tables[0]

	for row in [lista[i:i + 2] for i in range(0, len(lista), 2)]:
		#[suministro,B/A,fecha,activo/inactivo,numero de ordenesj,mono/tri,tarifa,ordenepre,departamento]
		try:
			basico = row[0]
			alternativo = row[1]
			if basico[3] == '2':
				print('El punto {} básico está dado de baja'.format(puntero))
				error = True
			if alternativo[3] == '2':
				print('El punto {} alternativo está dado de baja'.format(puntero))
				error = True
			plantilla.tables[0].column_cells(1)[puntero].text = basico[0]
			plantilla.tables[0].column_cells(2)[puntero].text = alternativo[0]
			plantilla.tables[0].column_cells(3)[puntero].text = basico[-1]
			tipo = int(basico[6])
			tipoAlternativo = int(alternativo[6])
			if tipo != tipoAlternativo:
				print('El punto {} tiene basico y alternativo, uno monofasico y otro trifasico'.format(puntero))
				error = True
			if tipo == 1:
				tablaDePuntos.column_cells(4)[puntero].text = 'Monofásica'
			elif tipo == 3:
				tablaDePuntos.column_cells(4)[puntero].text = 'Trifásica'
			else: print(tipo)
		except: pass
		finally:
			puntero+=1
	for i in range(17):
		for j in range(5):
			tablaDePuntos.column_cells(j)[i].paragraphs[0].style.font.size = Pt(12)
			tablaDePuntos.column_cells(j)[i].paragraphs[0].alignment = 1
	for i in range(17-puntero):
		tablaDePuntos._tbl.remove(tablaDePuntos.rows[puntero]._tr)

	#FORMATO
	plantilla.paragraphs[13].text = plantilla.paragraphs[13].text.replace('dd/mm/yy',sub('(?<=\d\d/\d\d/)\d\d(?=\d\d)','',fecha))
	if puntero == 2:
		plantilla.paragraphs[13].text = plantilla.paragraphs[13].text.replace('de los [cantidad] puntos en los','del punto en el')
	else:
		try: plantilla.paragraphs[13].text = plantilla.paragraphs[13].text.replace('[cantidad]','{} ({})'.format(cantidad[str(puntero-1)],puntero-1))
		except:
			print('No se encontraron mediciones para',fecha)
			print('1. Puede que estén mal las fechas que ingresaste en la tabla\n2.La ruta de la DB sea incorrecta. Fijarse en el archivo rutas.py.')
			input()
			exit()
		
	plantilla.paragraphs[13].style.font.size = Pt(12)
	plantilla.paragraphs[13].style.font.bold = 0

	#GUARDADO
	filename = 'Nota Selección de Puntos {} - {} - {} ({}).docx'.format(fecha[:2],fecha[3:5],fecha[-4:],lista[0][-1])
	while True:
		try:
			plantilla.save('Notas Seleccion de Puntos/{}'.format(filename))
			break
		except PermissionError:
			input('El archivo de la selección está abierto. Cerralo y apretá ENTER para volver a intentarlo')
			continue

	if error: 
		input()
		exit()
	else:
		subprocess.Popen('"{}" "Notas Seleccion de Puntos/{}"'.format(wordExe,filename),shell=True)

if __name__ == '__main__': main()