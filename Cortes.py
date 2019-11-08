import pyodbc
import subprocess
from docx import Document
from docx.shared import Pt
from re import sub,search
from time import mktime,strptime,strftime,localtime,time
from datetime import datetime
from dateutil.relativedelta import relativedelta
from rutas import dbCortes,dbUsuarios
from wordExe import wordExe
from Informacion import cortesEspeciales,choices,fechaFinRegistrosCortes

#VARIABLES
periodoDeAnalisis = 6 # meses
FinRegistrosCortes = datetime.strptime(fechaFinRegistrosCortes,'%d/%m/%y')
InicioRegistrosCortes = FinRegistrosCortes - relativedelta(months=periodoDeAnalisis,days=-1)
tablaUsuarios = 'suministros_sidac'
reCortesNoPenalizables = 'FOT|FOD|FIU|FCL|CSC'
printResultados = True

def getTablas():
	connCortes = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbCortes)
	cursorCortes = connCortes.cursor()
	tablas = [x[2] for x in cursorCortes.tables() if not x[2].startswith('MS')]
	tablas = sorted(tablas,key = lambda x: int(search('^\d{1,2}',x).group()))
	tablaCortes = tablas[-1]
	tablaCortesAnterior = tablas[-2]
	return tablaCortes,tablaCortesAnterior
	
def fechasCortes():
	global InicioRegistrosCortes,FinRegistrosCortes
	if InicioRegistrosCortes.day > 25:
		mList = ['0','1','2','3','4','5','6','7','8','9','10','11','12','1']
		mes = mList[InicioRegistrosCortes.month+1]
		año = str(InicioRegistrosCortes.year) if mes != '1' else str(InicioRegistrosCortes.year+1)
		dia = '01'
		fecha = '/'.join([dia,mes,año[2:]])
		InicioRegistrosCortes = datetime.strptime(fecha,'%d/%m/%y')
	if FinRegistrosCortes.month != InicioRegistrosCortes.month or __name__ != '__main__':
		return [datetime.strftime(InicioRegistrosCortes,'%d/%m/%y'),datetime.strftime(FinRegistrosCortes,'%d/%m/%y')]
	else: return False

def analisisCortes(cortes,tipoSuministro):
	penaliza = False
	if printResultados: print()
	cortes = limpiarCortes(cortes)
	cantidad = len(cortes)
	if 'T3' in tipoSuministro or 'T4' in tipoSuministro:
		if 'AT' in tipoSuministro:
			C,D = 3,120
		elif 'MT' in tipoSuministro:
			C,D = 4,180
		elif 'BT' in tipoSuministro:
			C,D = 6,360
	else:
		C,D = 6,600

	if cantidad > C:
		if printResultados: print('Penaliza por cantidad de cortes (>{}): {}'.format(C,round(cantidad)))
		penaliza = True
	tiempos = 0
	for corte in cortes:
		duracionRaw = round((corte[4]-corte[3]).total_seconds()/60,10)
		tiempos += duracionRaw
	if round(tiempos) > D:
		if printResultados: print('Penaliza por duracion de cortes (>{}): {}'.format(D,round(tiempos)))
		penaliza = True
	if not penaliza and printResultados: print('No penaliza: {} Cortes; {} Minutos'.format(round(cantidad),round(tiempos)))
	if printResultados: print('Tarifa:',tipoSuministro)
	return penaliza#?

def limpiarCortes(cortes):
	cortesPenalizables = []
	for corte in cortes:
		duracionRaw = round((corte[4]-corte[3]).total_seconds()/60,10)
		if duracionRaw <= 3 or search(reCortesNoPenalizables,corte[6]):	continue
		cortesPenalizables.append(corte)
			
	return cortesPenalizables

def procesarCortes(plantilla=None):
	fechas = fechasCortes()
	connCortes = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbCortes)
	cursorCortes = connCortes.cursor()
	if fechas:
		query = sub('TABLA2',tablaCortesAnterior,sub('TABLA1',tablaCortes,'SELECT [TABLA1].Suministro, [TABLA1].Interrupcion, [TABLA1].Orden_Reposicion, [TABLA1].Inicio, [TABLA1].Final, [TABLA1].Id, [TABLA1].Motivo_EPRE FROM [TABLA1] WHERE (Suministro = ? and Inicio <= DATEVALUE(?)) UNION SELECT [TABLA2].Suministro, [TABLA2].Interrupcion, [TABLA2].Orden_Reposicion, [TABLA2].Inicio, [TABLA2].Final, [TABLA2].Id, [TABLA2].Motivo_EPRE FROM [TABLA2] WHERE (Suministro = ? and Inicio >= DATEVALUE(?));'))
		cursorCortes.execute(query,(suministro,fechas[1],suministro,fechas[0]))
		cortes = cursorCortes.fetchall()
		cortes = sorted(cortes, key = lambda x: (x[3] - x[3].utcfromtimestamp(0)).total_seconds())
	else:
		query = sub('TABLA',tablaCortesAnterior,'SELECT [TABLA].Suministro, [TABLA].Interrupcion, [TABLA].Orden_Reposicion, [TABLA].Inicio, [TABLA].Final, [TABLA].Id, [TABLA].Motivo_EPRE FROM [TABLA] WHERE Suministro = ?;')
		cursorCortes.execute(query,(suministro))
		cortes = cursorCortes.fetchall()
		cortes = sorted(cortes, key = lambda x: (x[3] - x[3].utcfromtimestamp(0)).total_seconds())
	connUsuarios = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbUsuarios)
	cursorUsuarios = connUsuarios.cursor()
	query = 'SELECT {}.TARIFA,{}.CLIENTE FROM {} WHERE CLIENTE = ?'.format(tablaUsuarios,tablaUsuarios,tablaUsuarios)
	cursorUsuarios.execute(query,(suministro))
	res = cursorUsuarios.fetchall()
	if res:
		tipoSuministro = res[0][0]
	else:
		print('No está la tarifa en la tabla de suministros_sidac')
		if __name__ == '__main__':
			input()
			exit()
		else: return False,False,False,None
		
	if not cortes:
		print('El suministro {} no tuvo cortes desde el {} al {}'.format(suministro,fechas[0],fechas[1]))
		if __name__ != '__main__': return fechas,False,False,tipoSuministro
	else:
		penalizaCortes = analisisCortes(cortes,tipoSuministro)
		if __name__ != '__main__':
			return fechas,cortes,penalizaCortes,tipoSuministro
		plantilla.paragraphs[0].text = 'Fechas: {} al {}'.format(fechas[0],fechas[1])
		fila = 1
		for corte in cortes:
			plantilla.tables[0].rows[fila].cells[0].text = corte[1]
			plantilla.tables[0].rows[fila].cells[1].text = corte[3].strftime('%d/%m/%Y %H:%M')
			plantilla.tables[0].rows[fila].cells[2].text = corte[4].strftime('%d/%m/%Y %H:%M')
			duracion = round((corte[4]-corte[3]).total_seconds()/60)
			plantilla.tables[0].rows[fila].cells[3].text = str(duracion)
			plantilla.tables[0].rows[fila].cells[4].text = corte[6]
			c=0
			for _ in plantilla.tables[0].rows[fila].cells:
				for run in plantilla.tables[0].rows[fila].cells[c].paragraphs[0].runs:
					font = run.font
					font.size = Pt(11)
				plantilla.tables[0].rows[fila].cells[c].paragraphs[0].alignment = 1
				c+=1
			plantilla.tables[0].add_row()
			fila+=1
		plantilla.tables[0]._tbl.remove(plantilla.tables[0].rows[fila]._tr)
		plantilla.tables[0].style = 'Table Grid'
	return fechas

def forImport(sum,inicio=None,fin=None,printRes=True):
	global suministro,printResultados,FinRegistrosCortes,InicioRegistrosCortes,tablaCortes,tablaCortesAnterior
	tablaCortes,tablaCortesAnterior = getTablas()
	printResultados = printRes
	suministro = sum
	if fin:	FinRegistrosCortes = datetime.strptime(fin,'%d/%m/%y')
	if inicio: InicioRegistrosCortes = datetime.strptime(inicio,'%d/%m/%y')
	return procesarCortes()

if __name__ == '__main__':
	tablaCortes,tablaCortesAnterior = getTablas()
	
	print('Ver cortes del suministro:\n')
	print('XXXXXXXXXXX\r',end='')
	suministro = input()

	plantilla = Document('Recursos/Cortes.docx')
	fechas = procesarCortes(plantilla=plantilla)

	try:
		plantilla.save('Recursos/Cprocesado.docx')
	except:
		print('El archivo de cortes está abierto. Cerralo y apretá Enter')
		input()
		plantilla.save('Recursos/Cprocesado.docx')

	choice = input('\n¿Abrir tabla? s/N> ')
	if choice and choice.lower() in choices['yes']:
		subprocess.Popen('"{}" Recursos/Cprocesado.docx'.format(wordExe),shell = True)

	print('\nListo...')
	input()

