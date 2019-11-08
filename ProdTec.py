import pyodbc
import openpyxl
import win32api
import win32print
from Informacion import *
from re import search,sub
from time import strptime, strftime, sleep, localtime, mktime
from os import walk, getcwd, listdir
from os.path import isdir,isfile
from sys import argv
from docx import Document
from docx.shared import Pt
from functools import reduce
from openpyxl.chart import LineChart,Reference
from openpyxl.chart.text import RichText
from openpyxl.drawing.text import Paragraph, ParagraphProperties, CharacterProperties, Font
from openpyxl.worksheet import Worksheet
from openpyxl.reader.worksheet import WorkSheetParser
from openpyxl.worksheet.merge import MergeCells
from openpyxl.worksheet.cell_range import CellRange
from openpyxl.utils import range_boundaries
from rutas import plantillaCMP,plantillaSMP,db,dbHisto,unitariosDoc,planillaCompactos

DEBUG = False

class Registro():
	def __init__(self,fases,reg,dat):
		reg = reg.strip().replace(',','.').split('\t')
		self.penaliza = False
		self.fecha = reg[0]
		self.horario = reg[1]
		self.hora = int(reg[1].split(':')[0])
		self.V1 = float(reg[2])
		self.V1max = float(reg[3])
		self.V1min = float(reg[4])
		if fases == 3:
			try:
				self.V2 = float(reg[8])
				self.V2max = float(reg[9])
				self.V2min = float(reg[10])
				self.V3 = float(reg[14])
				self.V3max = float(reg[15])
				self.V3min = float(reg[16])
				self.thd = float(reg[20])
				self.flicker = float(reg[21])
				self.energia = float(reg[23])
				self.Vmax = max([self.V1max,self.V2max,self.V3max])
				self.Vmin = max([self.V1min,self.V2min,self.V3min])
				try: self.anormalidad = True if reg[24] == 'A' else None
				except: self.anormalidad = False
				self.corrupto = False
			except (IndexError,ValueError):
				print('EL dat {} tiene menos columnas que las necesarias'.format(dat))
				if __name__ == '__main__':salir()
				else: self.corrupto = True
		elif fases == 1:
			self.thd = float(reg[5])
			self.flicker = float(reg[6])
			try: self.anormalidad = True if reg[7] == 'A' else None
			except: self.anormalidad = False
			self.Vmax = self.V1max
			self.Vmin = self.V1min
			self.corrupto = False
		self.desequilibrio = None

class Medicion():
	def __init__(self,sum,dataMedicion,dataSuministro):
		self.suministro 	= sum
		self.dat 			= dataMedicion['dat']
		self.r32			= dataMedicion['r32']
		self.energia 		= dataMedicion['energia']
		self.fechaMedicion 	= dataMedicion['fecha']
		self.isMed			= True if dataMedicion['isMed'] == 'O' else False
		self.departamento 	= dataSuministro['departamento']
		self.usuario 		= dataSuministro['usuario']
		self.direccion 		= dataSuministro['direccion']
		self.tarifa 		= dataSuministro['tarifa']
		self.clase				= tipoTarifas[self.tarifa]
		self.tensionNominal	= tensionesNominales[self.clase]
		self.categoria 		= categorias[dataSuministro['categoria']]
		self.tipo 			= dataSuministro['tipo']
		self.seta 			= dataSuministro['seta']
		self.curvaDeCarga	= Ki[relacionTarifas[self.tarifa]]
		self.unitarios		= valoresUnitarios[dataSuministro['categoria']]
		# Limites
		self.lInf			= self.tensionNominal*valoresLimite[self.clase]['baja']
		self.lSup			= self.tensionNominal*valoresLimite[self.clase]['alta']
		self.lSupPen 		= self.tensionNominal*(1+valoresLimite[dataSuministro['categoria']][self.clase])
		self.lInfPen 		= self.tensionNominal*(1-valoresLimite[dataSuministro['categoria']][self.clase])
		self.lFlicker 		= valoresLimite['flicker']
		self.lThd 			= valoresLimite['thd'][self.clase]
		self.maximoRegistrosFueraDeRango = valoresLimite['Maximo porcentaje de registros fuera de rango']
		# Registros
		self.registros		= []
		
	def calcularMulta(self,totalRegistros,registrosPenalizados):
		multaTotal =			0
		multaUp = 				0
		multaDown = 			0
		energiaPenalizadaUp = 	0
		energiaPenalizadaDown = 0
		if DEBUG:
			debugDump = dumpDebug(self.dat)
			debugDump.send(None)
		for reg in registrosPenalizados:
			if self.fases == 3:
				v,desvioPorcentual = max(map(lambda x: (x,abs(x-self.tensionNominal)/self.tensionNominal),[reg.V1,reg.V2,reg.V3]),key= lambda x:x[1])# Desvios porcentuales
				difTensiones = v-self.tensionNominal
			elif self.fases == 1:
				difTensiones = reg.V1-self.tensionNominal
				desvioPorcentual = abs(difTensiones/self.tensionNominal)
			
			desviaciones = sorted(self.unitarios.keys())+[1]
			for desviacion in desviaciones:
				if desvioPorcentual > desviacion:
					precioKilowatt = self.unitarios[desviacion]
				else: break
			else: precioKilowatt = 0.0
			
			if self.fases == 3:
				multaParcial = reg.energia*precioKilowatt
				multaTotal += multaParcial
				if difTensiones > 0:
					multaUp += multaParcial
					energiaPenalizadaUp += reg.energia
				else:
					multaDown += multaParcial
					energiaPenalizadaDown += reg.energia
			elif self.fases == 1:
				multaParcial = (self.energia/totalRegistros)*self.curvaDeCarga[reg.hora]*precioKilowatt
				multaTotal += multaParcial
				if difTensiones > 0:
					multaUp += multaParcial
					energiaPenalizadaUp += self.energia/totalRegistros*self.curvaDeCarga[reg.hora]
				else:
					multaDown += multaParcial
					energiaPenalizadaDown += self.energia/totalRegistros*self.curvaDeCarga[reg.hora]

			if DEBUG:
				debugDump.send((self.fases,reg,desvioPorcentual,precioKilowatt,multaParcial,))
			
		self.multaTotal = multaTotal
		self.multaUp = multaUp
		self.multaDown = multaDown
		self.energiaPenalizadaUp = energiaPenalizadaUp
		self.energiaPenalizadaDown = energiaPenalizadaDown
			
	def getDat(self):
		"""
		devuelve un diccionario con la informacion que se encuentra en el archivo .dat
		es llamada por main()
		"""
		def errorDat(mensaje='No especificado. Ver archivo.'):
			print('El dat {} está corrupto posiblemente por una medición fallida\nError: {}'.format(search('\w+\.dat',f.name).group(),mensaje))
			
		
		try:
			with open(dir+'/'+self.dat,'r') as f:
				readFile = f.readlines()
		except UnicodeDecodeError:
			with open(dir+'/'+self.dat,'r',encoding='latin-1') as f:
				readFile = f.readlines()
			
		fases = 1 if len(readFile[7].split('\t')) == 8 else 3
		if fases == 3: self.energia = 0
		
		try:
			TV = float(search('(?<=Factor de Corrección: )\d+,\d+',readFile[2].split('\t')[3].strip().lstrip()).group().replace(',','.'))
			TI = float(search('(?<=Factor de Corrección: )\d+,\d+',readFile[3].split('\t')[3].strip().lstrip()).group().replace(',','.'))
		except:
			errorDat(mensaje='No se encontró el factor de tensión o corriente en la celda D3')
			return
		
		### Esto calcula el periodo. Si hay algun problema volver a la forma vieja ###
		## FORMA VIEJA: 
		# try: periodo,unidad = readFile[1].split('\t')[1].strip().lstrip().split()
		# except:	periodo,unidad = readFile[1].split('\t')[-1].strip().lstrip().split()[1:]
		## FORMA NUEVA:
		paraIntervalo = sub('\s+',' ',readFile[1]).split()
		for id,i in enumerate(paraIntervalo):
			if 'intervalo' in i.lower() or 'periodo' in i.lower() or '::' in i:
				periodo,unidad = paraIntervalo[id+1],paraIntervalo[id+2]
				break
		periodo = int(periodo)*60 if unidad == 'min.' else int(periodo) if unidad == 'seg.' else None

		for reg in readFile[9:]:
			registro = Registro(fases,reg,self.dat)
			if registro.corrupto: return False
			self.registros.append(registro)
			
		tmp = self.registros[::]
		if len(list(filter(lambda x: not x.anormalidad,self.registros))) == 0:
			print('CUIDADO: El archivo tiene todos los registros marcados como ANORMALIDADES.')
			for reg in tmp: reg.anormalidad = False
			self.registros = tmp
		else:		
			anormalidadPrevia = False
			anormalidadPosterior = False
			primero = True
				
						
			tmp = self.registros[::]
			for reg in tmp:
				if reg.anormalidad and primero:
					try: anormalidadPosterior = self.registros[self.registros.index(reg)+1].anormalidad
					except: pass
					if self.registros.index(reg) != 0: anormalidadPrevia = self.registros[self.registros.index(reg)-1].anormalidad

					if anormalidadPosterior: self.registros.remove(reg)
					else:
						self.registros.remove(reg)
						primero = False
						
			anormalidadPrevia = False
			anormalidadPosterior = False
			primero = True
			self.registros.reverse()
			tmp = self.registros[::]
			for reg in tmp:
				if reg.anormalidad and primero:				
					try: anormalidadPosterior = self.registros[self.registros.index(reg)+1].anormalidad
					except: pass
					if self.registros.index(reg) != 0: anormalidadPrevia = self.registros[self.registros.index(reg)-1].anormalidad

					if anormalidadPosterior: self.registros.remove(reg)
					else: primero = False
			self.registros.reverse()
		
			tmp = self.registros[::]
			for reg in tmp:
				if fases == 3:
					if not all([reg.V1,reg.V2,reg.V3,reg.V1max,reg.V2max,reg.V3max,reg.V1min,reg.V2min,reg.V3min,]) and reg.anormalidad:
						self.registros.remove(reg)
				elif fases == 1:
					if not all([reg.V1,reg.V1max,reg.V1min]) and reg.anormalidad:
						self.registros.remove(reg)

		tmp = self.registros[::]
		for registro in tmp:
			try:
				if fases == 3:
					if not inRange([registro.V1,registro.V2,registro.V3],self.lInf,self.lSup):
						self.registros.remove(registro)
				elif fases == 1:
					if not inRange([registro.V1],self.lInf,self.lSup):
						self.registros.remove(registro)
				primero = False
			except ValueError: continue
			
		if len(self.registros) == 0:
			if self.clase in ['MT13,2','MT33']:
				if self.suministro in compactos:
					factores = '\nTV: {}; TI: {}'.format(compactos[self.suministro]['tv'],compactos[self.suministro]['ti'])
				else: factores = ''
				print('Reprocesar el archivo R32 como MT.'.format(self.dat),factores)
			else: print('Verificar en base de datos si la tarifa del suministro corresponde a MT')
			return False
		elif len(self.registros) < 432000/periodo:
			print('El archivo tiene menos registros que los necesarios para el periodo de medición.')
			self.fallida = True

		try:
			fechaInicio = strftime('%d/%m/%Y',strptime(self.registros[0].fecha,'%d/%m/%y'))
			fechaFin = strftime('%d/%m/%Y',strptime(self.registros[-1].fecha,'%d/%m/%y'))
		except ValueError:
			fechaInicio = strftime('%d/%m/%Y',strptime(self.registros[0].fecha,'%d/%m/%Y'))
			fechaFin = strftime('%d/%m/%Y',strptime(self.registros[-1].fecha,'%d/%m/%Y'))
		if unidad == 'seg.':
			horaInicio = strftime('%I:%M:%S %p',strptime(self.registros[0].horario,'%H:%M:%S'))
			horaFin = strftime('%I:%M:%S %p',strptime(self.registros[-1].horario,'%H:%M:%S'))
		else:
			horaInicio = strftime('%I:%M:%S %p',strptime(self.registros[0].horario,'%H:%M'))
			horaFin = strftime('%I:%M:%S %p',strptime(self.registros[-1].horario,'%H:%M'))
		
		self.fases 			= fases
		self.periodo 		= periodo
		self.TV				= TV
		self.TI				= TI
		self.horaInicio 	= horaInicio
		self.fechaInicio 	= fechaInicio
		self.horaFin 		= horaFin
		self.fechaFin		= fechaFin
		return True
			
	def procesar(self):
		registrosCortes = []
		registrosCorrectos,registrosPenalizados = [],[]
		registrosSubTension,registrosSobreTension = 0,0
		registrosSubTensionF1,registrosSubTensionF2,registrosSubTensionF3 = 0,0,0
		registrosSobreTensionF1,registrosSobreTensionF2,registrosSobreTensionF3 = 0,0,0
		energiasPenalizadasSobre,energiasPenalizadasSub = 0,0
		registrosSubTensionCalc,registrosSobreTensionCalc = [],[]
		flickerFueraDeRango,thdFueraDeRango = [],[]
		desequilibrios = []
		
		if self.fases == 3:
			self.promedioVF1,self.promedioVF2,self.promedioVF3 = map(promedio,[[reg.V1 for reg in self.registros],[reg.V2 for reg in self.registros],[reg.V3 for reg in self.registros]])
			self.promedioVtotal = promedio([self.promedioVF1,self.promedioVF2,self.promedioVF3])
			for reg in self.registros:
				minimo,maximo = min(reg.V1,reg.V2,reg.V3),max(reg.V1,reg.V2,reg.V3)
				if minimo < self.lInf or maximo > self.lSup:
					registrosCortes.append(reg)
					continue
					
				self.energia += reg.energia
				
				registrosSubTensionF1 = registrosSubTensionF1+1 if reg.V1 < self.lInfPen else registrosSubTensionF1
				registrosSubTensionF2 = registrosSubTensionF2+1 if reg.V2 < self.lInfPen else registrosSubTensionF2
				registrosSubTensionF3 = registrosSubTensionF3+1 if reg.V3 < self.lInfPen else registrosSubTensionF3
				registrosSobreTensionF1 = registrosSobreTensionF1+1 if reg.V1 > self.lSupPen else registrosSobreTensionF1
				registrosSobreTensionF2 = registrosSobreTensionF2+1 if reg.V2 > self.lSupPen else registrosSobreTensionF2
				registrosSobreTensionF3 = registrosSobreTensionF3+1 if reg.V3 > self.lSupPen else registrosSobreTensionF3

				if minimo < self.lInfPen:
					registrosSubTension += 1
					energiasPenalizadasSub += reg.energia
					reg.penaliza = True
				if maximo > self.lSupPen:
					registrosSobreTension += 1
					energiasPenalizadasSobre += reg.energia
					reg.penaliza = True
				if not reg.penaliza: registrosCorrectos.append(reg)
				else: registrosPenalizados.append(reg)
					
				d1,d2,d3 = map(lambda x: abs(x-self.tensionNominal)/self.tensionNominal,[reg.V1,reg.V2,reg.V3])
				desequilibrios.append([d1,d2,d3])
				
				if reg.flicker > self.lFlicker:
					flickerFueraDeRango.append(reg.flicker)
				if reg.thd > self.lThd:
					thdFueraDeRango.append(reg.thd)
				
		elif self.fases == 1:
			self.promedioVtotal = promedio([reg.V1 for reg in self.registros])
			for reg in self.registros:
				if reg.V1 < self.lInf or reg.V1 > self.lSup:
					registrosCortes.append(reg)
					continue
				if reg.V1 < self.lInfPen:
					registrosSubTension += 1
					registrosSubTensionF1 += 1
					reg.penaliza = True
				if reg.V1 > self.lSupPen:
					registrosSobreTension += 1
					registrosSobreTensionF1 += 1
					reg.penaliza = True
				if not reg.penaliza: registrosCorrectos.append(reg)
				else: registrosPenalizados.append(reg)
				
				desequilibrios.append([abs(reg.V1-self.tensionNominal)/self.tensionNominal])
					
				if reg.flicker > self.lFlicker:
					flickerFueraDeRango.append(reg)
				if reg.thd > self.lThd:
					thdFueraDeRango.append(reg)
		
		registros = registrosPenalizados+registrosCorrectos
		totalRegistros = len(registros)
		totalRegistrosPenalizados = registrosSobreTension+registrosSubTension
		
		# calculo de multa	
		self.calcularMulta(totalRegistros,registrosPenalizados)

		self.penaliza = None
		if 'fallida' not in self.__dict__ :
			try:
				if totalRegistrosPenalizados/totalRegistros > self.maximoRegistrosFueraDeRango:
					self.penaliza = True
				else:
					self.penaliza = False
					self.multaDown = 0
					self.multaUp = 0
					self.multaTotal = 0
				self.fallida = None
			except ZeroDivisionError:
				print('El dat {} es posible fue procesado como BT y es posible que sea de MT o AT. Reprocesar el archivo R32'.format(self.dat))
				self.fallida = True
		
		self.penalizaFlicker = True if len(flickerFueraDeRango) > totalRegistros*0.05 else False
		self.penalizaThd = True if len(thdFueraDeRango) > totalRegistros*0.05 else False
			
		DF1 = [x[0] for x in desequilibrios]
		if self.fases == 3:
			DF2,DF3 = [x[1] for x in desequilibrios],[x[2] for x in desequilibrios]
			promediosDesequilibrios = list(map(promedio,[DF1,DF2,DF3]))
			try: desequilibriosMaximos = list(map(max,[DF1,DF2,DF3]))
			except ValueError: desequilibriosMaximos = [0,0,0]
		elif self.fases == 1:
			promediosDesequilibrios = promedio(DF1)
			try: desequilibriosMaximos = max(DF1)
			except ValueError: desequilibriosMaximos = [0]
		
		self.apartamientoMaximo = desequilibriosMaximos*100 if self.fases == 1 else max(desequilibriosMaximos)
		self.apartamientoMaximoF1 = desequilibriosMaximos*100 if self.fases == 1 else max(desequilibriosMaximos)
		self.apartamientoPromedio = promediosDesequilibrios*100 if self.fases == 1 else max(promediosDesequilibrios)
		self.apartamientoPromedioF1 = promediosDesequilibrios*100 if self.fases == 1 else promediosDesequilibrios[0]
		self.archivo = self.r32
		self.año = str(fecha.tm_year)
		self.distribuidora = distribuidoras[distribuidora]
		self.energiaPenalizada = self.energiaPenalizadaUp+self.energiaPenalizadaDown
		self.energiaSobretension = self.energiaPenalizadaUp
		self.energiaSubtension = self.energiaPenalizadaDown
		self.energiaTotal = self.energia
		self.energiaTotalF1 = self.energia
		self.flicker = promedio([x.flicker for x in self.registros])
		self.flickerF1 = self.flicker
		self.flickerFueraDeRango = len(flickerFueraDeRango)
		self.flickerFueraDeRangoF1 = len(flickerFueraDeRango)
		self.flickerFueraDeRangoF2, self.flickerFueraDeRangoF3 = '0,000','0,000'
		self.flickerPenalizable = 'SI' if self.penalizaFlicker else 'NO'
		self.flickerPenalizableF1 = 'SI' if self.penalizaFlicker else 'NO'
		self.mes = [x for x in meses if int(meses[x]) == fecha.tm_mon][0].capitalize()
		self.multaFueraDeRango = self.multaTotal
		self.promedioTension = self.promedioVtotal
		self.promedioTensionF1 = self.promedioVtotal if self.fases == 1 else self.promedioVF1
		self.promediosDesequilibrios = promediosDesequilibrios
		self.resultado = 'Penalizada' if self.penaliza else 'No Penalizada' if not self.fallida else 'Fallida'
		self.tensionMaxima = max([x.Vmax for x in self.registros])
		self.tensionMaximaF1 = self.tensionMaxima if self.fases == 1 else max([x.V1max for x in self.registros])
		self.tensionMinima = min([x.Vmin for x in self.registros])
		self.tensionMinimaF1 = self.tensionMinima if self.fases == 1 else min([x.V1min for x in self.registros])
		self.thdF1 = promedio([x.thd for x in self.registros])
		self.thdF2 = '-'
		self.thdF3 = '-'
		self.thdFueraDeRango = len(thdFueraDeRango)
		self.thdFueraDeRangoF1 = len(thdFueraDeRango)
		self.thdPenalizable = 'SI' if self.penalizaThd else 'NO'
		self.thdPenalizableF1 = self.thdPenalizable
		self.thdTotal = self.thdF1
		self.tipo = tipoInstalacion[self.fases]
		self.tipoMedicion = 'CME-O' if self.fases == 3 else 'SME-O'
		self.tipoMulta = 'Ambos' if (self.multaUp and self.multaDown) else 'Sub' if self.multaDown else 'Sobre' if self.multaUp else ''
		self.totalRegistros = totalRegistros
		self.totalRegistrosF1 = totalRegistros
		self.totalRegistrosPenalizados = totalRegistrosPenalizados
		self.totalRegistrosPenalizadosF1 = registrosSobreTensionF1+registrosSubTensionF1
		self.totalRegistrosSobretension = registrosSobreTension
		self.totalRegistrosSobretensionF1 = registrosSobreTensionF1
		self.totalRegistrosSubtension = registrosSubTension
		self.totalRegistrosSubtensionF1 = registrosSubTensionF1
		self.flickerPenalizableF2,self.flickerPenalizableF3 = 'NO','NO'
		self.promedioTensionF2 = self.promedioVF2 										if self.fases == 3 else '-'
		self.promedioTensionF3 = self.promedioVF3 										if self.fases == 3 else '-'
		self.tensionMaximaF2 = max([x.V2max for x in self.registros]) 					if self.fases == 3 else '-'
		self.tensionMaximaF3 = max([x.V3max for x in self.registros]) 					if self.fases == 3 else '-'
		self.tensionMinimaF2 = min([x.V2min for x in self.registros]) 					if self.fases == 3 else '-'
		self.tensionMinimaF3 = min([x.V3min for x in self.registros])					if self.fases == 3 else '-'
		self.thdPenalizableF2,self.thdPenalizableF3 = 'NO','NO'
		self.thdFueraDeRangoF2,self.thdFueraDeRangoF3 = '0,000','0,000'
		self.totalRegistrosF2 = totalRegistros 											if self.fases == 3 else '-'
		self.totalRegistrosF3 = totalRegistros 											if self.fases == 3 else '-'
		self.apartamientoMaximoF2 = desequilibriosMaximos[1] 							if self.fases == 3 else '-'
		self.apartamientoMaximoF3 = desequilibriosMaximos[2] 							if self.fases == 3 else '-'
		self.apartamientoPromedioF2 = promediosDesequilibrios[1] 						if self.fases == 3 else '-'
		self.apartamientoPromedioF3 = promediosDesequilibrios[2] 						if self.fases == 3 else '-'
		self.totalRegistrosPenalizadosF2 = registrosSobreTensionF2+registrosSubTensionF2 if self.fases == 3 else '-'
		self.totalRegistrosPenalizadosF3 = registrosSobreTensionF3+registrosSubTensionF3 if self.fases == 3 else '-'
		self.totalRegistrosSobretensionF2 = registrosSobreTensionF2 					if self.fases == 3 else '-'
		self.totalRegistrosSobretensionF3 = registrosSobreTensionF3 					if self.fases == 3 else '-'
		self.totalRegistrosSubtensionF2 = registrosSubTensionF2 						if self.fases == 3 else '-'
		self.totalRegistrosSubtensionF3 = registrosSubTensionF3 						if self.fases == 3 else '-'

def patch_worksheet():
    def merge_cells(self, range_string=None, start_row=None, start_column=None, end_row=None, end_column=None):
        cr = CellRange(range_string=range_string, min_col=start_column, min_row=start_row,
                      max_col=end_column, max_row=end_row)
        self.merged_cells.add(cr.coord)
    Worksheet.merge_cells = merge_cells
    def parse_merge(self, element):
        merged = MergeCells.from_tree(element)
        self.ws.merged_cells.ranges = merged.mergeCell
    WorkSheetParser.parse_merge = parse_merge
	
def dumpDebug(archivo):
	with open('{}/{}Debug.csv'.format(dir.replace('\\','/'),archivo),'w') as debug:
		header = ';'.join(('dia','hora','v1','v2','v3','desvio %','unitario','energia','multa parcial'))
		debug.write(header+'\n')
		while True:
			chunk = []
			fases,reg,desvioPorcentual,precioKilowatt,multaParcial = yield
			chunk.append(reg.fecha)
			chunk.append(reg.horario)
			if fases == 3:
				chunk.append(reg.V1)
				chunk.append(reg.V2)
				chunk.append(reg.V3)
			else:
				chunk.append(reg.V1)
				chunk.append('')
				chunk.append('')
			chunk.append(desvioPorcentual)
			chunk.append(precioKilowatt)
			if fases == 3:
				chunk.append(reg.energia)
			else:
				chunk.append('')
			chunk.append(multaParcial)
			debug.write(';'.join((str(x) for x in chunk)).replace('.',',')+'\n')
			
def inRange(numeros,lInf,lSup):
	"""
	Se fija si cada valor de la lista "numeros" está dentro de los rangos que se le entregan a la funcion.
	Si todos los valores están dentro de todos los rangos, devuelve True
	Si algún valor está por fuera de algún rango, devuelve False
	"""
	for numero in numeros:
		if numero < lSup and numero > lInf: continue
		else: break
	else: return True
	return False
	
def promedio(numeros):
	try:
		promedio = float(sum(numeros))/len(numeros)
	except ZeroDivisionError:
		promedio = 0.0
	return promedio
	
def salir(mensaje=''):
	"""
	Imprime el mensaje que se le entregue como argumento y termina el programa.
	"""
	print(mensaje,'\nEnter para terminar...')
	input()
	exit()
		
def getDataSuministros(suministros):
	"""
	devuelve un diccionario con la informacion los suministros que estan en la tabla Usuarios en la db.
	{suministro:{diccionario con informacion},...}
	"""
	error = []
	sumsDict = {}
	# Se conecta con la db
	conn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %db)
	cursor = conn.cursor()
	# Arma la consulta
	sumsToQuery = ','.join(["'%s'"%x for x in suministros])
	listaDatos = ['suministro','departamento','usuario','direccion','tarifa','tipo','seta','categoria']
	if distribuidora == 'ESJ':
		query = 'SELECT [tablaUsuarios].[NUMERO_DE_SUMINISTRO], [tablaUsuarios].[DISTRITO], [tablaUsuarios].[NOMBRE_ APELLIDO], [tablaUsuarios].[DIRECCION], [tablaUsuarios].[TARIFA], [tablaUsuarios].[TIPO], [tablaUsuarios].[NUMERO_DE_CENTRO], [tablaUsuarios].[categoria] FROM [tablaUsuarios] WHERE [tablaUsuarios].[NUMERO_DE_SUMINISTRO] in ({});'
	elif distribuidora == 'DECSA':
		query = 'SELECT [tablaUsuarios].[NUMERO_DE_SUMINISTRO], [tablaUsuarios].[DISTRITO], [tablaUsuarios].[NOMBRE_ APELLIDO], [tablaUsuarios].[DIRECCION], [tablaUsuarios].[TARIFA], [tablaUsuarios].[TIPO], [tablaUsuarios].[NUMERO_DE_CENTRO], [tablaUsuarios].[categoria] FROM [tablaUsuarios] WHERE [tablaUsuarios].[NUMERO_DE_SUMINISTRO] in ({});'.format(sumsToQuery)
	cursor.execute(query.format(sumsToQuery).replace('tablaUsuarios',tablasUsuarios[distribuidora]),())
	# Ejecuta la consulta
	lista = cursor.fetchall()
	# Se fija si estan todos los datos. Si no encuentra alguno te avisa
	for dato in lista:
		dataSuministro = dict(zip(listaDatos,dato))
		for key in dataSuministro:
			if dataSuministro[key] == None:
				print(' '.join(['"',key.upper(),'"','no esta especificada en la tabla Usuarios para',dataSuministro['suministro']]))
				if key.lower() in ['direccion',]:
					valor = input('Ingresar {} (ENTER para saltear) :> '.format(key))
					dataSuministro[key] = valor
				else: error.append(dataSuministro['suministro'])
		sumsDict[str(dataSuministro['suministro'])] = dataSuministro
	if error:
		print('¿Continuar con los que tienen todos los datos?\nS/n> \r',end='S/n> ')
		while True:
			choice = input().lower()
			if choice in choices['yes']:
				for sum in error:
					sumsDict[str(sum)] = False
				break
			elif choice in choices['no']:
				salir()
			else:
				print('"s" o "n"\nS/n> \r',end='S/n> ')
		error = False
	for sum in suministros:
		if str(sum) not in sumsDict.keys():
			error = True
			print('El suministro {} no se encuentra en la base de datos'.format(sum))
			if distribuidora == 'DECSA':
				sumsDict[str(sum)] = {'suministro':str(sum),'departamento':'Caucete','usuario':'-','direccion':'-','tarifa':'T1-R1','tipo':'3','seta':str(sum),'categoria':'AE'}
				
	if error: 
		if distribuidora == 'DECSA':
			print('Se autocompletan los faltantes como suministros aereos y tarifa T1-R1\n')
		else:
			print('Revise la tabla de usuarios de {}'.format(distribuidora))
			if __name__ == '__main__':salir()
			else: return False
	
	#Si esta todo bien devuelve el diccionario
	return sumsDict
	
def getTablaMediciones():
	pulmonDeDias = 2
	conn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %db)
	cursor = conn.cursor()
	query = 'SELECT * FROM [Seleccion de Puntos] WHERE DATEVALUE(?)<=[Seleccion de Puntos].[Fecha Instalacion] AND [Seleccion de Puntos].[Fecha Instalacion]<=DATEVALUE(?);'
	fechaPrevia = strftime('%d/%m/%y',localtime(mktime(fecha)-(60*60*24*pulmonDeDias)))
	fechaPosterior = strftime('%d/%m/%y',localtime(mktime(fecha)+(60*60*24*pulmonDeDias)))
	cursor.execute(query,(fechaPrevia,fechaPosterior))
	tabla = cursor.fetchall()
	mediciones = dict([(x[1],(x[0],)+x[2:]) for x in tabla])
	return mediciones
	
def getExpedientesRemedicion(suministro):
	conn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %db)
	cursor = conn.cursor()
	query = 'SELECT [Expediente] FROM [Seleccion de puntos] WHERE ([Seleccion de puntos].[Pto de medición] = ? AND [Expediente] <> Null) GROUP BY [Expediente]'
	cursor.execute(query,(suministro,))
	listaExpedientes = cursor.fetchall()
	if listaExpedientes:
		listaExpedientes = [x[0] for x in listaExpedientes]
	return listaExpedientes

def getDataMediciones(resultados,fecha,dats):
	"""
	devuelve un diccionario con la siguiente estructura:
	{'12345678912':{'dat':'11111111.dat','energia':56.8}, ...} para ESJ
	{'0057':{'dat':'2478818O10000000.dat','energia':0.0}, ...} para DECSA
	"""
	error = False
	suministros = {}
	resultadosMediciones = openpyxl.load_workbook('{}/{}'.format(dir,resultados))[sheets[distribuidora]]
	if distribuidora == 'ESJ':
		fecha = strftime('%d/%m/%Y',fecha)
		for dat in dats:
			fila = 9
			while resultadosMediciones['A%d'%fila].value != None:
				try: fRes = resultadosMediciones['F%d'%fila].value.split(' ')[0]
				except AttributeError: fRes = strftime('%d/%m/%Y',(resultadosMediciones['F%d'%fila].value.timetuple()))
				if fecha == fRes:
					try:
						r32 = resultadosMediciones['Q%d'%fila].value
						datName = r32.split('.')[0]+'.dat'
					except AttributeError:
						fila += 1
						continue
					if datName == dat:
						energia = float(resultadosMediciones['N%d'%fila].value)-float(resultadosMediciones['K%d'%fila].value)
						isMed = resultadosMediciones['C%d'%fila].value
						suministros[str(resultadosMediciones['A%d'%fila].value)] = {'dat':datName,'energia':energia,'fecha':fecha,'r32':r32,'isMed':isMed}
						break
				fila += 1
			else:
				print('No se encontró {} en el informe de {} para el {}'.format(dat,distribuidora,fecha))
				error = True
		if error:
			if __name__ == '__main__': salir()
			else: return False
	elif distribuidora == 'DECSA':
		for dat in dats:
			fila = 5
			while resultadosMediciones['B%d'%fila].value != None:
				try: fechaExcel = str(resultadosMediciones['I%d'%fila].value.strftime('%d/%m/%Y'))
				except:
					if search('(?i)fallida',resultadosMediciones['H%d'%fila].value):
						fila+=1
						continue
					else: salir('Faltan fechas en el informe de DECSA')
				r32 = resultadosMediciones['B%d'%fila].value+'.R32'
				datName = resultadosMediciones['B%d'%fila].value+'.dat'
				if datName == dat:
					energia = 0.0
					suministros[str(resultadosMediciones['C%d'%fila].value).zfill(4)] = {'dat':datName,'r32':r32,'energia':energia,'fecha':fechaExcel,'isMed':'O'}
					break
				fila += 1
			else:
				print('No se encontró {} en el informe de la distribuidora'.format(dat))
				error = True
		if error: salir()
	return suministros

def armarExcel(medicion):
	"""
	Genera el excel resumen del procesamiento de cada medicion y lo guarda en el mismo directorio
	donde están los dats.
	"""
	
	fases = medicion.fases
	plantilla = plantillaCMP if fases == 3 else plantillaSMP
	celdas = celdasPlantillaTrifasica if fases == 3 else celdasPlantillaMonofasica
	wb = openpyxl.load_workbook(plantilla)
	ws = wb['Plantilla']

	for dato in celdas:
		if dato in medicion.__dict__:
			ws[celdas[dato]].value = medicion.__dict__[dato]
		
	tensiones = [[]]+[[reg.V1] for reg in medicion.registros]
	armarGrafico(fases,1,tensiones,ws,medicion)
	if fases == 3:
		tensiones = [[]]+[[reg.V2] for reg in medicion.registros]
		armarGrafico(fases,2,tensiones,ws,medicion)
		tensiones = [[]]+[[reg.V3] for reg in medicion.registros]
		armarGrafico(fases,3,tensiones,ws,medicion)
	
	filename = dir+'/'+medicion.dat.split('.')[0]+'.xlsx'
	wb.save(filename)
	return filename
	
def armarGrafico(fases,fase,tensiones,ws,medicion):
	"""
	Arma los gráficos de excel que van a ir en los archivos
	"""

	posicion = posicionesGraficos[fases][fase]
	c=1
	for V in tensiones[1:]:
		ws.cell(row=c, column=17+fase, value=V[0])
		c+=1
		
	###### Dibuja los limites superior e inferior
	c=1
	for _ in tensiones[1:]:
		ws.cell(row=c, column=22, value=medicion.lSupPen)
		c+=1
	c=1
	for _ in tensiones[1:]:
		ws.cell(row=c, column=23, value=medicion.lInfPen)
		c+=1
	######

	c1 = LineChart()
	c1.title = "Fase {}".format(fase)
	c1.style = 13
	c1.y_axis.title = 'Tensión [V]'
	c1.x_axis.title = 'Registros'

	vList = [x[0] for x in tensiones[1:]]
	minimo = min([min(vList),medicion.lInfPen])
	maximo = max([max(vList),medicion.lSupPen])
	c1.y_axis.scaling.min = minimo*0.95
	c1.y_axis.scaling.max = maximo*1.05
	
	font = Font(typeface='Verdana')
	size = 1250
	cp = CharacterProperties(latin=font, sz=size, b=False)
	c1.x_axis.title.tx.rich.p[0].pPr = ParagraphProperties(defRPr=cp)
	c1.y_axis.title.tx.rich.p[0].pPr = ParagraphProperties(defRPr=cp)
	c1.title.tx.rich.p[0].pPr = ParagraphProperties(defRPr=cp)
	c1.x_axis.txPr = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=cp), endParaRPr=cp)])
	c1.y_axis.txPr = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=cp), endParaRPr=cp)])
	c1.title.txPr = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=cp), endParaRPr=cp)])
	c1.x_axis.tickLblSkip = 200
	c1.x_axis.scaling.min = 0
	c1.legend = None

	data = Reference(ws, min_col=17+fase, min_row=1, max_col=17+fase, max_row=len(tensiones))
	limiteSuperior = Reference(ws, min_col=22, min_row=1, max_col=22, max_row=len(tensiones))
	limiteInferior = Reference(ws, min_col=23, min_row=1, max_col=23, max_row=len(tensiones))
	c1.add_data(data, titles_from_data=True)
	c1.add_data(limiteSuperior, titles_from_data=True)
	c1.add_data(limiteInferior, titles_from_data=True)

	s2 = c1.series[0]
	s2.graphicalProperties.line.width = Pt(1)
	s2.graphicalProperties.line.solidFill = "000000"
	c1.series[1].graphicalProperties.line.width = Pt(2.2)
	c1.series[1].graphicalProperties.line.solidFill = "000000"
	c1.series[2].graphicalProperties.line.width = Pt(2.2)
	c1.series[2].graphicalProperties.line.solidFill = "000000"

	c1.height = 12
	c1.width = 24
	ws.add_chart(c1, posicion)

def getFecha():
	"""
	devuelve una fecha a partir de la carpeta donde se esta trabajando.
	Si la carpeta no tiene el dia en su nombre, se pregunta por la fecha para insertarla manualmente
	Si la distribuidora es DECSA, pregunta por el mes de medicion
	es llamada por main()
	"""
	global dia
	meses = {'enero':'01','febrero':'02','marzo':'03','abril':'04','mayo':'05','junio':'06','julio':'07','agosto':'08','septiembre':'09','octubre':'10','noviembre':'11','diciembre':'12'}
	
	try:
		año = search('20\d\d',dir).group()
	except:
		año = str(localtime().tm_year)
		if '-y' not in argv:
			print('Año de las mediciones:> {}\r'.format(año),end='Año de las mediciones:> ')
			año = input()
		else: año = argv[argv.index('-y')+1]
		if not año: año = str(localtime().tm_year)
		print()
	print('Usando año {}'.format(año))
	
	if distribuidora == 'DECSA':
		for mes in meses:
			if mes in dir.lower():
				fecha = mes
				break
		else:
			print('Mes de medicion> \r',end = 'Mes de medicion> ')
			fecha = input()
		dia = mes
		fecha = strptime('01/{}/{}'.format(meses[fecha.lower()],año[-2:]),'%d/%m/%y')
	elif distribuidora == 'ESJ':
		diaLocal = search('(?i)\d?\d de \w{4,10}',dir)
		if diaLocal:
			dia = diaLocal.group()
			diaLocal = diaLocal.group().lower()
			for mes in meses:
				if mes in diaLocal:
					diaLocal = diaLocal.replace(mes,meses[mes])+' '+año
			fecha = strptime(diaLocal,'%d de %m %Y')
		else:
			while True:
				print('Fecha> xx/xx/xxxx\r',end = 'Fecha> ')
				diaLocal = input()
				try:
					fecha = strptime(diaLocal,'%d/%m/%Y')
					dia = '{} de {}'.format(str(fecha.tm_mday).zfill(2),[x for x in meses if meses[x] == str(fecha.tm_mon).zfill(2)][0])
					break
				except ValueError:
					print('Fecha mal formateada. Ejemplo: 20/09/'+año)
	return fecha,año

def cargarUnitarios(factorDeInversion):
	"""
	Devuelve un diccionario con la siguiente estructura:
	{'AE':{DESVIACION:VALOR_UNITARIO,DESVIACION:VALOR_UNITARIO,...},
	{'SU':....},
	{'RU':....},
	}
	"""
	global valoresUnitarios
	for doc in reversed(unitariosDoc):
		try: stampStr = search('\d{1,2}-\d{1,2}-\d{4}',doc).group()
		except: continue
		stamp = strptime(stampStr,'%d-%m-%Y')
		if mktime(fecha) > mktime(stamp):
			docFile = Document(doc)
			print("""
			Usando unitarios vigentes desde el {}
			Usando factor de inversion = {}
			""".format(strftime('%d/%m/%Y',stamp),factorDeInversion))
			break
	else: salir(mensaje='\nNo se encontraron los unitarios para {}'.format(strftime('%d/%m/%Y',fecha)))

	valoresUnitarios = {}
	rangos = [.05,.06,.07,.08,.09,.1,.11,.12,.13,.14,.15,.16,.18]
	valoresUnitarios['SU'] = dict(zip(rangos,[float(x.text.replace(',','.'))*factorDeInversion for x in docFile.tables[0].columns[3].cells[1:]]))
	valoresUnitarios['AE'] = dict(zip(rangos[2:],[float(x.text.replace(',','.'))*factorDeInversion for x in docFile.tables[1].columns[3].cells[1:]]))
	valoresUnitarios['RU'] = dict(zip(rangos[2:],[float(x.text.replace(',','.'))*factorDeInversion for x in docFile.tables[2].columns[3].cells[1:]]))

def getFiles():
	"""
	devuelve un diccionario con los dats y el resultado que manda la distribuidora
	Declara un global con la distribuidora que es usado por varias funciones
	es llamada por main()
	"""
	global distribuidora
	tmp = []
	resultadosESJ = ''
	resultadosDECSA = ''

	for root, folder, file in walk(dir):
		tmp.append(file)
	files = [x for x in tmp[0]]
	dats = [x for x in files if x[-4:] == '.dat']
	for f in files:
		if search('(?i)Resultados (?:Re)?mediciones',f) and f.endswith('.xlsx'):
			resultadosESJ = f
			distribuidora = 'ESJ'
			break
		elif search('TABLA RESULTADO (?:RE)?MEDICIONES -',f) and f.endswith('.xlsx'):
			resultadosDECSA = f
			distribuidora = 'DECSA'
			break
	else:
		print('No se encontro el informe de resultados de la distribuidora. Verificá la extensión = xlsx')
		if __name__ == '__main__': salir()
		else: return False
	return {'dats':dats,'resultados':resultadosESJ if resultadosESJ else resultadosDECSA if resultadosDECSA else None}

def getDir(folder):
	"""
	Define un global con el Directorio donde se va a trabajar.
	No devuelve nada porque "dir" es global.
	Si no se le pone la direccion, el programa pone como defecto el directorio
	donde se encuentra a la hora de ejecución.
	"""
	global dir
	if folder: dir = folder
	else:
		while True:
			dir = input('Carpeta:> ')
			if not isdir(r'{}'.format(dir)):
				print('No existe el directorio ingresado.')
				continue
			else: break
		if not dir: dir = getcwd()
	
def cargarHisto(mediciones):
	"""
	Carga los resultados de las mediciones procesadas en la tabla histo
	"""	
	lista = []
	tabla = 'Histo' if distribuidora == 'ESJ' else 'HistoDECSA' if distribuidora == 'DECSA' else None
	for medicion in mediciones:
		t = []
		tmp = (medicion.suministro, dia, medicion.año, medicion.tipoMulta, medicion.multaUp, medicion.multaDown, medicion.multaFueraDeRango, medicion.tipoMedicion, medicion.fechaInicio, medicion.horaInicio, medicion.fechaFin, medicion.horaFin, medicion.totalRegistros, medicion.totalRegistrosF1, medicion.totalRegistrosF2, medicion.totalRegistrosF3, medicion.totalRegistrosSobretension, medicion.totalRegistrosSobretensionF1, medicion.totalRegistrosSobretensionF2, medicion.totalRegistrosSobretensionF3, medicion.totalRegistrosSubtension, medicion.totalRegistrosSubtensionF1, medicion.totalRegistrosSubtensionF2, medicion.totalRegistrosSubtensionF3, medicion.totalRegistrosPenalizados, medicion.totalRegistrosPenalizadosF1, medicion.totalRegistrosPenalizadosF2, medicion.totalRegistrosPenalizadosF3, medicion.energiaTotal, medicion.energiaPenalizadaUp, medicion.energiaPenalizadaDown, medicion.promedioTension, medicion.promedioTensionF1, medicion.promedioTensionF2, medicion.promedioTensionF3, medicion.tensionMaxima, medicion.tensionMaximaF1, medicion.tensionMaximaF2, medicion.tensionMaximaF3, medicion.tensionMinima, medicion.tensionMinimaF1, medicion.tensionMinimaF2, medicion.tensionMinimaF3, medicion.thdFueraDeRango, medicion.flickerFueraDeRango, medicion.resultado, dir, medicion.archivo)
		for i in tmp:
			if i == '-': t.append('0')
			else: t.append(str(i))
		lista.append(tuple(t))

	conn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbHisto)
	cursor = conn.cursor()
	query = 'INSERT INTO {}(Suministro,Mes,Año,Tension,Multa_Sobre,Multa_Sub,Multa,Tipo_Medicion,Fecha_Colocacion,Hora_Colocacion,Fecha_Retiro,Hora_Retiro,Reg_Procesados,Reg_Procesados_Fase1,Reg_Procesados_Fase2,Reg_Procesados_Fase3,Reg_Sobre_Total,Reg_Sobre_Fase1,Reg_Sobre_Fase2,Reg_Sobre_Fase3,Reg_Sub_Total,Reg_Sub_Fase1,Reg_Sub_Fase2,Reg_Sub_Fase3,Reg_Penal_Total,Reg_Penal_Fase1,Reg_Penal_Fase2,Reg_Penal_Fase3,Energía_Total,Energía_Sobre,Energía_Sub,V_Prom_Total,V_Prom_Fase1,V_Prom_Fase2,V_Prom_Fase3,V_Max_Total,V_Max_Fase1,V_Max_Fase2,V_Max_Fase3,V_Min_Total,V_Min_Fase1,V_Min_Fase2,V_Min_Fase3,Armonicas_Fuera,Flicker_Fuera,Resultado,Archivo_Destino,Archivo) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?);'.format(tabla)
	cursor.executemany(query,lista)
	cursor.commit()
	
def cargarCompactos():
	global compactos
	wb = openpyxl.load_workbook(planillaCompactos)
	ws = wb['compactos']
	compactos = {}
	for row in ws.rows:
		sum = row[0].value
		if not sum.isnumeric(): continue
		try:
			relTvp = float(row[8].value)
			relTvs = float(row[9].value)
			tv = relTvp/relTvs
			relTip = float(row[6].value)
			relTis = float(row[7].value)
			ti = relTip/relTis
		except TypeError:
			continue
		compactos[sum] = {'tv':tv,'ti':ti}
	return compactos
	
	
def imprimir(archivosExcel,choice):		
	if choice in [1,2]:
		for medicion in archivosExcel:
			if medicion['penaliza']: # Si penaliza...
				win32api.ShellExecute(0,"print",medicion['archivo'],None,".",0)
				sleep(2)
			if choice != 2:
				win32api.ShellExecute(0,"print",medicion['archivo'],None,".",0)
				sleep(2)		
	else: return


def checkRutasDB(rutas):
	problem = False
	for id,ruta in enumerate(rutas,start=1):
		largoSeñalador = 0
		c = 0
		rutaSliced = ruta.split('/')
		for index,slice in enumerate(rutaSliced[:-1]):
			path = reduce(lambda x,y:'/'.join((x,y)), rutaSliced[:index+1])
			if c <= 2:
				c += 1
				largoSeñalador += len(slice)+1
				continue
			if not isdir(path):
				print(str(index)+')',ruta)
				señalador = ''.join(('   ',' '*(largoSeñalador),' '*(int(len(slice)/2)),'^'))
				señalador += ''.join(('\n___','_'*(largoSeñalador),'_'*(int(len(slice)/2)),'|'))
				print(señalador)
				print('\t','La carpeta no existe')
				problem = True
				break
			largoSeñalador += len(slice)+1
			c += 1
		else:
			if not isfile(ruta):
				archivos = listdir('/'.join(rutaSliced[:-1]))
				if not rutaSliced[-1].split('.')[-1]:
					print(str(id)+')',ruta)
					print('\t','El archivo no tiene extensión')
					problem = True
					continue
				filename = rutaSliced[-1].split('.')[0]
				similares = [x for x in archivos if filename in x]
				if not similares:
					print(str(id)+')',ruta)
					señalador = ''.join(('   ',' '*(len('/'.join(rutaSliced[:-1]))),' '*(int(len(rutaSliced[-1].split('.')[0])/2)),'^'))
					señalador += ''.join(('\n___','_'*(len('/'.join(rutaSliced[:-1]))),'_'*(int(len(rutaSliced[-1].split('.')[0])/2)),'|'))
					print(señalador)
					print('\t','El nombre del archivo está mal o el fichero no existe')
					problem = True
				else:
					print(str(id)+')',ruta)
					señalador = ''.join(('   ',' '*(len(ruta)-2),'^'))
					señalador += ''.join(('\n___','_'*(len(ruta)-2),'|'))
					print(señalador)
					print('\t','La extension es incorrecta')
					problem = True		
	if problem:
		salir()

@manejarErrores
def main(folder=None,histo=True,imp=True):
	checkRutasDB((db,dbHisto))
	global fecha,año
	print()
	print('Usando db de usuarios: {}. Para modificarla, cambiar el archivo en rutas.py'.format(db))
	print()
	listaFinal = []
	archivosExcel = []
	mediciones = []
	getDir(folder)
	files = getFiles() # getFiles declara un global con la distribuidora. distribuidora = 'ESJ' | 'DECSA'
	if not files: return False
	if distribuidora == 'DECSA':
		import ArreglaDatsDECSA
		ArreglaDatsDECSA.arreglar(dir)
	fecha,año = getFecha() # timeStruct
	factorDeInversion = factoresDeInversion[distribuidora][sorted(list(filter(lambda x: mktime(fecha) > mktime(strptime(x,'%d/%m/%y')),factoresDeInversion[distribuidora])),key = lambda x: mktime(strptime(x,'%d/%m/%y')))[-1]]
	cargarUnitarios(factorDeInversion)
	cargarCompactos()
	dats = files['dats']
	resultados = files['resultados']
	tablaMediciones = getTablaMediciones()
	dataMediciones = getDataMediciones(resultados,fecha,dats) # {'12345678912':{'dat':'11111111.dat','energia':56.8},'fecha':'' ...} para ESJ
	if not dataMediciones: return False
	suministros = [str(x) for x in sorted(dataMediciones.keys(),key = lambda x: dataMediciones[x]['dat'][:2])]
	dataSuministros = getDataSuministros(suministros)
	if not dataSuministros: return False
	patch_worksheet()
	
	seleccion = None
	c = 1
	for sum in suministros:
		if not dataSuministros[sum]: continue
		medicion = Medicion(sum,dataMediciones[sum],dataSuministros[sum])
		print(medicion.dat.center(14,' ').center(50,'-'))
		if not medicion.getDat():
			print('Se saltea {}'.format(medicion.dat))
			continue
		medicion.procesar()
		excel = armarExcel(medicion)
		if sum in tablaMediciones:
			seleccion = tablaMediciones[sum]
		elif not medicion.isMed:
			expedientes = ','.join(getExpedientesRemedicion(sum))
			seleccion = (c,'','','','','REMEDICION',expedientes,c)
		elif medicion.isMed and distribuidora == 'ESJ':
			print('El suministro {} no fue pedido por el EPRE.'.format(sum))
			seleccion = (c,'','','','','PUNTO DESCONOCIDO','',-1)
		archivosExcel.append({'archivo':excel,'dat':medicion.r32,'penaliza':medicion.penaliza,'suministro':sum,'seleccion':seleccion})
		mediciones.append(medicion)
		print('Fallida' if medicion.fallida else 'Penaliza' if medicion.penaliza else 'No Penaliza')
		c+=1
	
	if seleccion: archivosExcel = sorted(archivosExcel,key=lambda x:x['seleccion'][0])
	else: print('\nNo se encontraron mediciones en la tabla de seleccion de puntos')
	
	if imp:
		print('\n¿Imprimir mediciones?\n1) SI\n2) Solo Penalizadas\n3) NO')
		while True:
			try:
				choice = int(input('> '))
				if choice not in [1,2,3]:
					print('Un valor entre 1 y 3')
					continue
				break
			except ValueError: continue
		exp = ''
		tmp = []
		if seleccion:
			for file in archivosExcel:
				if exp != file['seleccion'][5]:
					imprimir(tmp,choice)
					exp = file['seleccion'][5]
					if tmp != []: input('Continuar...')
					tmp = []
				tmp.append(file)
				expediente = file['seleccion'][5] if file['seleccion'][5] else ''
				print('\t',str(int(file['seleccion'][-1])).ljust(2),expediente.ljust(12),file['dat'],file['suministro'],file['seleccion'][6])
			if tmp: imprimir(tmp,choice)
		else: imprimir(archivosExcel,choice)
	print()
	if histo:
		if __name__ == '__main__':
			print('¿Guardar en la tabla Histo?')
			choice = input('S/n> ')
			if choice in choices['yes']: cargarHisto(mediciones)
			salir(mensaje='\nListo...')
		else:
			cargarHisto(mediciones)
	return True

if __name__ == '__main__': main()