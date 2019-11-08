import pyodbc
import subprocess
from docx import Document
from docx.shared import Pt
from re import sub,search
from time import mktime,strptime,strftime,localtime,time,gmtime
from rutas import dbMediciones
from wordExe import wordExe

#VARIABLES
reResultadoNoPenalizado = '(?i)correct(?:a|o)+|no penaliza(?:da)*|NP'
reResultadoFallido = '(?i)fa(?:ll)?(?:ida)?'
choices = {'yes':['y','s'],'no':['n','']}

def procesarMediciones(plant=None,suministros=None,setas=None):
	global plantilla
	
	connMediciones = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbMediciones)
	cursorMediciones = connMediciones.cursor()
	if suministros:
		query = 'SELECT [Base de Mediciones].[PTO-MED],[Base de Mediciones].[FECHA-COL], [Base de Mediciones].[FECHA- RET], [Base de Mediciones].[RESULTADO], [Base de Mediciones].[SETA] FROM [Base de Mediciones] WHERE [PTO-MED] IN ({}) ORDER BY [FECHA-COL] DESC;'.format(','.join(['?' for _ in suministros]))
		cursorMediciones.execute(query,(tuple(suministros)))
	elif setas:
		query = 'SELECT [Base de Mediciones].[PTO-MED],[Base de Mediciones].[FECHA-COL], [Base de Mediciones].[FECHA- RET], [Base de Mediciones].[RESULTADO], [Base de Mediciones].[SETA] FROM [Base de Mediciones] WHERE [SETA] IN ({}) ORDER BY [FECHA-COL] DESC;'.format(','.join(['?' for _ in setas]))
		cursorMediciones.execute(query,(tuple(setas)))
		
	mediciones = cursorMediciones.fetchall()
	
	if not mediciones:
		print('No se encontraron mediciones para los suministros en la base de datos')
		return False
	if suministros:
		medSumReclamo = [x[4] for x in mediciones if x[0] == suministros[0]]
		if not medSumReclamo:
			print('No se encontraron mediciones para el suministro ingresado. ¿Dado de baja?')

	if setas:
		suministros = list(set((x[0] for x in mediciones)))
	meds = {y:[x[1:] for x in mediciones if x[0] == y] for y in suministros}
	
	f = []
	for sum in suministros:
		if len(sum) != 11:
			print('Se saltea "{}".'.format(sum))
			continue
		temp = meds[sum]
		temp = sorted(temp, key = lambda x: (x[0] - x[0].utcfromtimestamp(0)).total_seconds())
		if not setas:
			print()
			print(sum)
			for i in range(len(temp)):
				res = temp[i][2].upper() if temp[i][2] else 'DESCONOCIDO'
				print('{}) {} {}'.format(i+1,temp[i][0].strftime('%d/%m/%Y'),res))
			while True:
				index = input(':> ')
				if not index: break
				index = int(index)
				if index > len(temp): print('Mmmm... ese número no está en la lista :(')
				else: break
			if not index: continue
		else:
			index = 1
		f+=[list(x)+[sum] for x in temp[index-1:]]
	meds = f
	if not meds and __name__ == '__main__': exit()
	meds = sorted(meds, key = lambda x: (x[0] - x[0].utcfromtimestamp(0)).total_seconds())
	
	if __name__ != '__main__': return meds

	plantilla = plant
	fila = 2
	for med in meds:
		plantilla.tables[0].rows[fila].cells[0].text = med[-1]
		plantilla.tables[0].rows[fila].cells[1].text = med[3] if med[3] else ''
		plantilla.tables[0].rows[fila].cells[2].text = med[0].strftime('%d/%m/%Y')
		plantilla.tables[0].rows[fila].cells[3].text = med[1].strftime('%d/%m/%Y')
		if search(reResultadoNoPenalizado,med[2]):
			plantilla.tables[0].rows[fila].cells[4].text = '"No Penalizada"'
		elif search(reResultadoFallido,med[2]):
			plantilla.tables[0].rows[fila].cells[4].text = '"Fallida"'
		else:
			plantilla.tables[0].rows[fila].cells[4].text = '"Penalizada"'
			
		c = 0
		for _ in plantilla.tables[0].rows[fila].cells:
			for run in plantilla.tables[0].rows[fila].cells[c].paragraphs[0].runs:
				font = run.font
				font.size = Pt(11)
			plantilla.tables[0].rows[fila].cells[c].paragraphs[0].alignment = 1
			c+=1
		for run in plantilla.tables[0].rows[fila].cells[4].paragraphs[0].runs:
				font = run.font
				font.italic = True
		plantilla.tables[0].add_row()
		fila+=1
	plantilla.tables[0]._tbl.remove(plantilla.tables[0].rows[fila]._tr)
	fechaInicioMedicion = plantilla.tables[0].rows[2].cells[2].text
	fechaFinMedicion = plantilla.tables[0].rows[2].cells[3].text
	plantilla.tables[0].style = 'Table Grid'
	return True

def forImport(sums):
	global suministrosMedidos
	suministrosMedidos = sums
	return procesarMediciones(suministros=suministrosMedidos)


if __name__ == '__main__':
	suministrosMedidos = []
	setasMedidas = []
	print('Suministro o seta:\n')
	while True:
		print('xxxxxxxxxxx\r',end='')
		sum = input()
		if len(sum) == 11:
			suministrosMedidos.append(sum)
		elif 0 < len(sum) <= 5:
			setasMedidas.append(sum)
		else: break

	plantilla = Document('Recursos/Mediciones.docx')
	if suministrosMedidos:
		procesarMediciones(plant=plantilla,suministros=suministrosMedidos)
	elif setasMedidas:
		procesarMediciones(plant=plantilla,setas=setasMedidas)
	else:
		input('Enter para terminar...')
		exit()
		

	plantilla.save('Recursos/Mprocesado.docx')

	choice = input('¿Abrir tabla? s/N> ')
	if choice.lower() in choices['yes']:
		subprocess.Popen('"{}" Recursos/Mprocesado.docx'.format(wordExe),shell = True)
	print('\nListo...')