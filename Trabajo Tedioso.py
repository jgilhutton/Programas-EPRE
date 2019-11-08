##################################################
##                                              ##
## Hecho por Juan Ignacio Gil-Hutton            ##
## Cualquier problema que encuentren            ##
## con el programa contáctenmennnnnn            ##
## a estos lugares:                             ##
## 	-Tel: 264 5067132                           ##
##  -mail: jgilhutton@gmail.com                 ##
## Hay una copia de todos los programas         ##
## en la siguiente dirección:                   ##
## https://github.com/jgilhutton/Programas-EPRE ##
##                                              ##
##################################################

from docx import Document
from docx.shared import Pt,RGBColor
from re import sub,search
from time import strftime,strptime,localtime
import Cortes
import CortesTresAños
import Mediciones
import DataSuministro
import pyodbc
from os import mkdir
from math import floor
from rutas import dbUsuarios
from Informacion import manejarErrores,choices,penalizaTipo,tiposUsuario,tiposSuministro,reResultadoFallido,reResultadoNoPenalizado,resultadosMediciones,plazos,valoresLimite,cantidad

class Informe:
	resultadosMedicionesParrafo = {'P':'la medición "penalizada" es aquella en la que no se han cumplido los límites para el nivel de tensión establecidos en el Contrato de Concesión','NP':'la medición "no penalizada" es aquella en la que no se han podido constatar incumplimientos a los estándares de nivel de tensión previstos en el Contrato de Concesión','F':'la medición "fallida" es aquella en la que no se ha alcanzado la cantidad mínima de registros prevista en el Contrato de Concesión'}
	parrafoAuditoria = 'Las acciones mencionadas por la Distribuidora han sido auditadas por el perito técnico del E.P.R.E., verificando la realización de las mismas.'
	parrafoFormulacionDeCargos = 'Posteriormente, ante la presunción de incumplimiento por parte de la Distribuidora, a lo dispuesto en la Resolución E.P.R.E. Nº {}, y simultáneamente se intimó a Energía San Juan S.A. a brindar solución a los inconvenientes relevados.'
	plantillas = {'IGGRES':'Recursos/IGG RES.docx',
				'IGG1NP':'Recursos/IGG 1NP.docx',
				'IGGFD':'Recursos/IGG Sumario Demoras.docx',
				'Resolucion':'Recursos/Resolucion.docx',
				'NU':'Recursos/NU.docx',
				'ND1NP':'Recursos/ND 1NP.docx',
				'IGG2P':'Recursos/IGG 2P.docx',
				'ND2P':'Recursos/ND 2P.docx',
				'Cortes3Años':'Recursos/Planilla Cortes.xlsx',
				'PaseArchivo':'Recursos/Pase Archivo.docx',
				'PaseArchivoResolucion':'Recursos/Pase Archivo Resolucion.docx',
				'CargosDemora':'Recursos/Sumario Demoras.docx',
				}
	limites = {'T1-R1': {'duracion': 36000.0, 'cantidad': 6.0},
				'T1-AP': {'duracion': 36000.0, 'cantidad': 6.0},
				'TRA-RCD': {'duracion': 21600.0, 'cantidad': 6.0},
				'T3-MT-13.2-B': {'duracion': 10800.0, 'cantidad': 4.0},
				'T1-G2': {'duracion': 36000.0, 'cantidad': 6.0},
				'T3-BT': {'duracion': 21600.0, 'cantidad': 6.0},
				'T1-R3': {'duracion': 36000.0, 'cantidad': 6.0},
				'TRA-RSD': {'duracion': 21600.0, 'cantidad': 6.0},
				'T4-MT-13.2-R': {'duracion': 10800.0, 'cantidad': 4.0},
				'T1-G1': {'duracion': 36000.0, 'cantidad': 6.0},
				'T2-SMP': {'duracion': 36000.0, 'cantidad': 6.0},
				'T3-MT-13.2-R': {'duracion': 10800.0, 'cantidad': 4.0},
				'T3-MT-33': {'duracion':10800.0, 'cantidad': 4.0},
				'TEDP3-BT': {'duracion': 21600.0, 'cantidad': 6.0},
				'T4-AT': {'duracion': 7200.0, 'cantidad': 3.0},
				'T1-R2': {'duracion': 36000.0, 'cantidad': 6.0},
				'T1R2-CONS': {'duracion': 36000.0, 'cantidad': 6.0},
				'T4-MT-33': {'duracion': 10800.0, 'cantidad': 4.0},
				'T1R1-CONS': {'duracion': 36000.0, 'cantidad': 6.0},
				'T4-BT': {'duracion': 21600.0, 'cantidad': 6.0},
				'T2-CMP': {'duracion': 36000.0, 'cantidad': 6.0},
				'TEDP2-CMP': {'duracion': 36000.0, 'cantidad':6.0},
				'TEDP-G': {'duracion': 36000.0, 'cantidad': 6.0},
				'T1R3-CONS': {'duracion': 36000.0, 'cantidad': 6.0},
				'T1-G3': {'duracion': 36000.0, 'cantidad': 6.0}}
		
	def __init__(self,parentDict,parentTipo=None):
		if parentDict:
			self.__dict__ =  parentDict
			self.amChild = True
		else:
			self.amChild = False
		self.parentTipo = parentTipo
	
	def mkdir(self,optPath = None):
		"""
		Crea el directorio donde se van a guardar los archivos generados.
		si "optPath" no es None, guarda en optPath
		"""
		if optPath:
			self.dir = 'Informes/{}'.format(optPath)
		else:
			self.dir = 'Informes/{} {}'.format(self.expedienteSup.replace('/','-'),self.usuarioSup)
		try:
			mkdir(self.dir)
		except FileExistsError:
			pass

	def openDocx(self,path):
		"""
		Abre el archivo .docx con direccion "path"
		"""
		documento = Document(path)
		return documento
		
	def getCortes3Años(self,sum,tarifa = False):
		"""
		Usa el modulo "CortesTresAños" para consultar los cortes.
		"""
		print('\tConsultando 3 años de cortes...')
		self.fechascortes3Años,self.cortes3Años,self.tarifa,self.seta = CortesTresAños.forImport(sum,tarifa)
		
		
	def getCortes(self):
		"""
		Usa el modulo "Cortes" para consultar los cortes.
		"""
		print('\tConsultando 6 meses de cortes...')
		self.fechasCortes,self.cortes,self.penalizaCortes,self.tarifa = Cortes.forImport(self.suministros[0])
		self.tipoPenalizacion(self.tarifa)
			
	def getMediciones(self):
		"""
		Usa el modulo "Mediciones" para consultar las mediciones.
		"""
		print('\tConsultando mediciones...')
		self.mediciones = Mediciones.forImport(self.suministrosMedidos)
		
	def tipoPenalizacion(self,tipoSuministro):
		"""
		Determina la tension nominal y los limites de cortes y tiempo a partir de la tarifa
		"""
		if 'T3' in tipoSuministro or 'T4' in tipoSuministro:
			if 'AT' in tipoSuministro:
				self.limCantCortes,self.limTiempoCortes = '3','2'
				self.tensionNominal = '33.000'
			elif 'MT' in tipoSuministro:
				self.limCantCortes,self.limTiempoCortes = '4','3'
				self.tensionNominal = '13.200'
			elif 'BT' in tipoSuministro:
				self.limCantCortes,self.limTiempoCortes = '6','6'
				self.tensionNominal = '220'
		else:
			self.limCantCortes,self.limTiempoCortes = '6','10'
			self.tensionNominal = '220'
			
	def completarTablaMediciones(self,plantilla,indexTablaMediciones):
		"""
		Toma las mediciones que devuelve "getMediciones()" y llena la tabla con indice "indexTablaMediciones" que se encuentra en "plantilla"
		"""
		if not self.mediciones: return plantilla
		print('\tCompletando tabla de mediciones...')
		fila = 2
		suministrosPenalizados = []
		for med in self.mediciones:
			plantilla.tables[indexTablaMediciones].rows[fila].cells[0].text = med[-1]
			plantilla.tables[indexTablaMediciones].rows[fila].cells[1].text = med[3] if med[3] else ""
			plantilla.tables[indexTablaMediciones].rows[fila].cells[2].text = med[0].strftime('%d/%m/%Y')
			plantilla.tables[indexTablaMediciones].rows[fila].cells[3].text = med[1].strftime('%d/%m/%Y')
			if search(reResultadoNoPenalizado,med[2]):
				plantilla.tables[indexTablaMediciones].rows[fila].cells[4].text = '"No Penalizada"'
				resultadosMediciones['NP'] = True
			elif search(reResultadoFallido,med[2]):
				plantilla.tables[indexTablaMediciones].rows[fila].cells[4].text = '"Fallida"'
				resultadosMediciones['F'] = True
			else:
				plantilla.tables[indexTablaMediciones].rows[fila].cells[4].text = '"Penalizada"'
				suministrosPenalizados.append(med)
				resultadosMediciones['P'] = True
			
			c = 0
			for _ in plantilla.tables[indexTablaMediciones].rows[fila].cells:
				for run in plantilla.tables[indexTablaMediciones].rows[fila].cells[c].paragraphs[0].runs:
					font = run.font
					font.size = Pt(11)
				plantilla.tables[indexTablaMediciones].rows[fila].cells[c].paragraphs[0].alignment = 1
				c+=1
			for run in plantilla.tables[indexTablaMediciones].rows[fila].cells[4].paragraphs[0].runs:
					font = run.font
					font.italic = True
			plantilla.tables[indexTablaMediciones].add_row()
			fila+=1
		plantilla.tables[indexTablaMediciones]._tbl.remove(plantilla.tables[indexTablaMediciones].rows[fila]._tr)
		self.fechaInicioMedicion = ''.join(search('(^\d{1,2}/\d{1,2}/).*((?<=20)\d\d$)',plantilla.tables[indexTablaMediciones].rows[-1].cells[2].text).groups())
		self.fechaFinMedicion = ''.join(search('(^\d{1,2}/\d{1,2}/).*((?<=20)\d\d$)',plantilla.tables[indexTablaMediciones].rows[-1].cells[3].text).groups())
		self.suministrosPenalizados = tuple([y[-1] for y in filter(lambda x: x[0].strftime('%d/%m/%y') == self.fechaInicioMedicion,suministrosPenalizados)])
		
		return plantilla
		
	def completarTablaCortes3Años(self,plantilla,indexTablaCortesTresAños,full=False):
		"""
		Toma los cortes que devuelve "getCortes3Años()" y llena la tabla con indice "indexTablaCortesTresAños" que se encuentra en "plantilla"
		Si "full" es True, llena la tabla con las columnas CodigoDeInterrupcion y MotivoEPRE
		"""
		print('\tCompletando tabla cortes 3 años...')
		fila = 1
		for corte in self.cortes3Años:
			# SUMINISTRO, INTERRUPCION, INICIO, FINAL, ID, MOTIVO EPRE
			if full: plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[0].text = corte[0]
			plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[0].text = corte[1]
			plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[1].text = corte[3].strftime('%d/%m/%Y %H:%M')
			plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[2].text = corte[4].strftime('%d/%m/%Y %H:%M')
			duracion = floor((corte[4]-corte[3]).total_seconds()/60)
			plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[3].text = str(duracion)
			if full: plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[5].text = corte[6]
			c=0
			for _ in plantilla.tables[indexTablaCortesTresAños].rows[fila].cells:
				for run in plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[c].paragraphs[0].runs:
					font = run.font
					font.size = Pt(11)
				plantilla.tables[indexTablaCortesTresAños].rows[fila].cells[c].paragraphs[0].alignment = 1
				c+=1
			plantilla.tables[indexTablaCortesTresAños].add_row()
			fila+=1
		plantilla.tables[indexTablaCortesTresAños]._tbl.remove(plantilla.tables[indexTablaCortesTresAños].rows[fila]._tr)
		
		return plantilla
		
	def completarTablaCortes(self,plantilla,indexTablaCortes):
		"""
		Toma los cortes que devuelve "getCortes()" y llena la tabla con indice "indexTablaCortes" que se encuentra en "plantilla"
		"""
		print('\tCompletando tabla de cortes...')
		fila = 1
		for corte in self.cortes:
			plantilla.tables[indexTablaCortes].rows[fila].cells[0].text = corte[1]
			plantilla.tables[indexTablaCortes].rows[fila].cells[1].text = corte[3].strftime('%d/%m/%Y %H:%M')
			plantilla.tables[indexTablaCortes].rows[fila].cells[2].text = corte[4].strftime('%d/%m/%Y %H:%M')
			duracion = floor((corte[4]-corte[3]).total_seconds()/60)
			plantilla.tables[indexTablaCortes].rows[fila].cells[3].text = str(duracion)
			c=0
			for _ in plantilla.tables[indexTablaCortes].rows[fila].cells:
				for run in plantilla.tables[indexTablaCortes].rows[fila].cells[c].paragraphs[0].runs:
					font = run.font
					font.size = Pt(11)
				plantilla.tables[indexTablaCortes].rows[fila].cells[c].paragraphs[0].alignment = 1
				c+=1
			plantilla.tables[indexTablaCortes].add_row()
			fila+=1
		plantilla.tables[indexTablaCortes]._tbl.remove(plantilla.tables[indexTablaCortes].rows[fila]._tr)
		
		return plantilla
	
	def replaceRuns(self,plantilla):
		"""
		Esta función se encarga de reemplazar todas las variables en "plantilla".
		Los archivos de plantillas deben tener las cosas a reeemplazar en azul(0000FF) y con el mismo nombre que las variables en este programa.
		"""
		vars = ['suministroReclamo','inicioTresAños','finTresAños','inicioCortes','finCortes']
		for parrafo in plantilla.paragraphs:
			for run in parrafo.runs:
				try:
					if str(run.element.rPr.color.val) == '0000FF':
						run.text = run.text.replace('{','')
						run.text = run.text.replace('}','')
						variable = search('(?<={)?\w{7,50}(?=})?',run.text).group()# Ese "{7,40}" puede generar problemas.
						if variable in self.__dict__ or variable in vars:
							if variable == 'suministroReclamo':
								run.text = sub(variable,', '.join(self.__dict__['suministros']),run.text)
							elif variable == 'inicioTresAños':
								run.text = sub(variable,self.fechascortes3Años[0],run.text)
							elif variable == 'finTresAños':
								run.text = sub(variable,self.fechascortes3Años[1],run.text)
							elif variable == 'inicioCortes':
								run.text = sub(variable,self.fechasCortes[0],run.text)
							elif variable == 'finCortes':
								run.text = sub(variable,self.fechasCortes[1],run.text)
							elif variable == 'semestresPenalizados':
								if not self.semestresPenalizados:
									run.text = sub(variable,'no',run.text)
								elif len(self.semestresPenalizados) == 1:
									run.text = sub(variable,'en el período correspondiente al '+self.semestresPenalizados[0],run.text)
								else:
									run.text = sub(variable,'en el período correspondiente al '+', '.join(self.semestresPenalizados[:-1])+' y '+self.semestresPenalizados[-1],run.text)
							else:
								run.text = run.text.replace(variable,self.__dict__[variable])
							if type(self) in [Nota25m,AvanceDeObra,DebidaRespuesta,NotaCPP]:
								run.element.rPr.color.val = RGBColor(0x00,0x00,0x00)
						else:
							pass
				except:
					pass
					
		return plantilla
	
	def borrarParrafo(self,parrafo):
		"""
		Hace falta explicar?
		"""
		p = parrafo._element
		p.getparent().remove(p)
		p._p = p._element = None
		
	def borrarTabla(self,tabla):
		"""
		Hace falta explicar?
		"""
		t = tabla._element
		t.getparent().remove(t)
		t._p = t._element = None
		
	def saveDocx(self,doc,path):
		"""
		Hace falta explicar?
		"""
		doc.save(path)
		
class Resolucion(Informe):
	"""
	Hace:
		informe a gerencia general
		proyecto de resolucion
		planilla de cortes 3 años
	"""
	def getData(self):
		if not self.amChild:
			print('Expediente:>')
			print('550.xxxx/xx\r',end='550.')
			e = input()
			self.expedienteSup = '550.' + e

			self.usuarioSup = input('Iniciador:> ')
			print('Tipo usuario:\n1) el\n2) la\n3)los\n4)las')
			self.tipoUsuario = tiposUsuario[int(input(':> '))]

			print('Reclamo:>')
			print('xxxx/xx\r',end='')
			self.reclamoSup = input()
			print('Direccion Postal')
			self.direccionPostal = input(':> ')
			print('Departamento')
			self.departamentoSup = input(':> ')

			print('Suministros de reclamo:\n')
			self.suministros = []
			while True:
				print('XXXXXXXXXXX\r',end='')
				sum = input()
				if len(sum) == 11: self.suministros.append(sum)
				elif sum == '': break
				elif sum == 'otros': self.suministros.append('otros')
				else:
					print('Suministro incorrecto')

			print('Fecha intimación:\n:> xx/xx/xxxx\r',end=':> ')		
			self.fechaIntimacion = input()

			print('Nota A.C.C.:\n')
			print('xxxx/xx\r',end='')
			self.notaDistribuidora = input()
			print('Fecha Nota A.C.C. #{}\n'.format(self.notaDistribuidora))
			print('xx/xx/xxxx\r',end='')
			self.fechaNotaDistribuidora = input()

			self.tipoSuministro = tiposSuministro[2] if len(self.suministros) > 1 else tiposSuministro[1]

			while True:
				self.accionesESJ = input('Acciones que tomó ESJ\n[...] se informó mediante nota {} que '.format(self.notaDistribuidora))
				confirmacion = input('ENTER para confirmar. Cualquier tecla+ENTER para volver a escribir')
				if not confirmacion: break
			self.suministrosMedidos = [self.suministros[0]]
			print('Suministros medidos:\n{}'.format(self.suministrosMedidos[0]))
			while True:
				print('xxxxxxxxxxx\r',end='')
				sum = input()
				if sum:
					self.suministrosMedidos.append(sum)
				else: break
			
			self.getMediciones()
			self.getCortes()
			self.resultadoCortesRes = 'estado dentro' if not self.penalizaCortes else 'sido apartada'
			self.apartamientoMaximo = str(int(valoresLimite[DataSuministro.query(self.suministros[0])['Categoria']]['BT']*100))
			self.apartamientoMaximoPalabra = cantidad[self.apartamientoMaximo]
			self.getCortes3Años(self.suministros[0],tarifa=True)

	def make(self):
		self.getData()
		if not self.amChild: self.mkdir()
		igg = GerenciaGeneral(self.__dict__,type(self))
		self.__dict__ = igg.make()
		excelCortes3A = ExcelCortes3A(self.__dict__)
		self.__dict__ = excelCortes3A.make()
		print('Generando proyecto de resolucion...')
		plantilla = self.openDocx(self.plantillas['Resolucion'])
		plantilla = self.completarTablaMediciones(plantilla,1)
		plantilla = self.completarTablaCortes(plantilla,0)
		plantilla = self.replaceRuns(plantilla)
		self.saveDocx(plantilla,'{}/{} {} - Resolucion.docx'.format(self.dir,self.expedienteSup.replace('/','-'),self.usuarioSup))
		
		return self.__dict__

class GerenciaGeneral(Informe):	
	"""
	Hace:
		informe a gerencia general
	"""
	def getData(self):
		if not self.amChild or self.parentTipo in [PrimeraNoPenalizada,SegundaPenalizada,PaseArchivo,CargosDemora]:
			print('Expediente:>')
			print('550.xxxx/xx\r',end='550.')
			e = input()
			self.expedienteSup = '550.' + e

			self.usuarioSup = input('Iniciador:> ')
			print('Tipo usuario:\n1) el\n2) la\n3)los\n4)las')
			self.tipoUsuario = tiposUsuario[int(input(':> '))]
			
			print('Suministros de reclamo:\n')
			self.suministros = []
			while True:
				print('XXXXXXXXXXX\r',end='')
				sum = input()
				if len(sum) == 11: self.suministros.append(sum)
				elif sum == '' and self.suministros: break
				elif sum == 'otros': self.suministros.append('otros')
				else:
					print('Suministro incorrecto')
					continue
				self.suministroReclamo = self.suministros[0]

			if self.parentTipo in [PaseArchivo,CargosDemora]:
				self.resolucionSup = input('Resolucion\n:> ')
				if self.resolucionSup:
					choice = int(input('Motivo de la resolucion:\n1) Tension\t2) Cortes\t3) Tension y Cortes\n:> '))
					if choice == 1:
						self.motivoResolucion = 'producto'
						self.motivoDetalle = 'Niveles de Tensión'
					elif choice == 2: 
						self.motivoResolucion = 'servicio'
						self.motivoDetalle = 'Cantidad y Duración de las Interrupciones'
					elif choice == 3:
						self.motivoResolucion = 'producto y servicio'
						self.motivoDetalle = 'Cantidad y Duración de las Interrupciones y Niveles de Tensión'
					if self.parentTipo == CargosDemora:
						self.fechaComunicacion = input('Fecha en la que fue comunicada la resolucion {}\n:> '.format(self.resolucionSup))

				
			if self.parentTipo in [SegundaPenalizada] or (self.parentTipo in [PaseArchivo] and not self.resolucionSup):
				print('Fechas de mediciones informadas previamente')
				print('Inicio Medicion:> xx/xx/xx\r',end='Inicio Medicion:> ')
				self.fechaInicioMedicionesPrevias = input()
				print('Fin Medicion:> xx/xx/xx\r',end='Fin Medicion:> ')
				self.fechaFinMedicionesPrevias = input()
				if self.parentTipo is not PaseArchivo:
					choice = input('1) Penalizadas\n2) No Penalizadas\n:> ')
					if choice == '1':
						self.resultadosMedicionesInformadas = 'apartados'
						self.detalleAcciones = 'brindar solución a los inconvenientes relevados'
					elif choice == '2':
						self.resultadosMedicionesInformadas = 'dentro'
						self.detalleAcciones = 'realizar tareas de normalización en caso de ser necesarias'
					
			
			if self.parentTipo == CargosDemora:
				notas = []
				print('Notas A.C.C. de obra:')
				while True:
					temp = input()
					if temp:
						notas.append(temp)
						print('xxxx/xx',end='\r')
					else: break
				if len(notas) == 1:
					self.obraPlSing = 'la Nota A.C.C. #'
					self.notaDistribuidora = notas[0]
				else:
					self.obraPlSing = 'las Notas A.C.C. #'
					self.notaDistribuidora = ', '.join(notas[:-1])+' y '+notas[-1]
			else:
				print('Nota A.C.C.:\n')
				print('xxxx/xx\r',end='')
				self.notaDistribuidora = input()
			print('Fecha Nota A.C.C. #{}\n'.format(self.notaDistribuidora))
			print('xx/xx/xx\r',end='')
			self.fechaNotaDistribuidora = input()

			self.tipoSuministro = tiposSuministro[2] if len(self.suministros) > 1 else tiposSuministro[1]

			while True:
				self.accionesESJ = input('Acciones que tomó ESJ\n[...] se informó mediante nota {} que '.format(self.notaDistribuidora))
				confirmacion = input('ENTER para confirmar. Cualquier tecla+ENTER para volver a escribir')
				if not confirmacion: break
				
			if self.parentTipo == PaseArchivo:
				self.accionesSumariales = True if input('¿Tiene formulación de cargos? S/n :> ').lower() in choices['yes'] else False
				if self.accionesSumariales:
					self.parrafoFormulacionDeCargos = self.parrafoFormulacionDeCargos.format(self.resolucionSup)
					print('Nota A.C.C.:\n')
					print('xxxx/xx\r',end='')
					self.notaDistribuidoraFormulacionDeCargos = input()
					print('Fecha Nota A.C.C. #{}\n'.format(self.notaDistribuidoraFormulacionDeCargos))
					print('xx/xx/xx\r',end='')
					self.fechaNotaFormulacionDeCargos = input()
					self.accionesEsjFormulacionDeCargos = input('Acciones que tomó ESJ\n[...] se informó mediante nota {} en fecha {} que '.format(self.notaDistribuidoraFormulacionDeCargos,self.fechaNotaFormulacionDeCargos))
				self.auditadoSiNo = self.parrafoAuditoria if input('¿Auditado? S/n:> ').lower() in choices['yes'] else ''
				
			if self.parentTipo != CargosDemora:
				self.suministrosMedidos = [self.suministros[0]]
				print('Suministros medidos:\n{}'.format(self.suministrosMedidos[0]))
				while True:
					print('xxxxxxxxxxx\r',end='')
					sum = input()
					if sum:
						self.suministrosMedidos.append(sum)
					else: break
			
			self.tipoSuministro = tiposSuministro[2] if len(self.suministros) > 1 else tiposSuministro[1]
			
			self.getCortes()
			try:
				dataSuministro = DataSuministro.query(self.suministros[0])
				if not dataSuministro:
					print('No se encontro el suministro en la tabla Usuarios. Estableciendo el apartamiento máximo como 3%')
					self.apartamientoMaximo = '3'
				else:
					self.apartamientoMaximo = str(int(valoresLimite[dataSuministro['Categoria']]['BT']*100))
			except KeyError:
				salir(mensaje = 'Hubo un error al intentar obtener los datos del suministro. Eliminá el archivo U.pickle de la carpeta Recursos y probá devuelta.')
			self.apartamientoMaximoPalabra = cantidad[self.apartamientoMaximo]
			if self.parentTipo == Resolucion:
				self.getCortes3Años(self.suministros[0])
			else: self.cortes3Años = None
			if self.parentTipo != CargosDemora:
				self.getMediciones()

			if not dataSuministro:
				for sum in self.suministros:
					dataSuministro = DataSuministro.query(sum)
					if dataSuministro: break
				else:
					if self.mediciones:
						for sum in [x[-1] for x in self.mediciones]:
							dataSuministro = DataSuministro.query(sum)
							if dataSuministro: break
			if not dataSuministro: self.subestacion = input('SETA:> ')
			else: self.subestacion = dataSuministro['Seta']
	
	def make(self):
		self.getData()
		if not self.amChild or self.parentTipo in [PrimeraNoPenalizada,SegundaPenalizada,PaseArchivo,CargosDemora]: self.mkdir()
		print('Generando informe a gerencia general...')
		if self.parentTipo == Resolucion:
			plantilla = self.openDocx(self.plantillas['IGGRES'])
		elif self.parentTipo == PrimeraNoPenalizada:
			plantilla = self.openDocx(self.plantillas['IGG1NP'])
		elif self.parentTipo == SegundaPenalizada:
			plantilla = self.openDocx(self.plantillas['IGG2P'])
		elif self.parentTipo == PaseArchivo:
			if self.resolucionSup:
				plantilla = self.openDocx(self.plantillas['PaseArchivoResolucion'])
			else:
				plantilla = self.openDocx(self.plantillas['PaseArchivo'])
		elif self.parentTipo == CargosDemora:
			plantilla = self.openDocx(self.plantillas['IGGFD'])
			
		if self.parentTipo != CargosDemora:
			plantilla = self.completarTablaMediciones(plantilla,0)
		if self.cortes and self.parentTipo != Resolucion:
			if self.parentTipo != CargosDemora:
				plantilla = self.completarTablaCortes(plantilla,1)
			else:
				plantilla = self.completarTablaCortes(plantilla,0)
		else:
			if self.parentTipo == PrimeraNoPenalizada:
				self.borrarParrafo(plantilla.paragraphs[21])
				self.borrarTabla(plantilla.tables[1])
				self.borrarParrafo(plantilla.paragraphs[22])
				self.borrarParrafo(plantilla.paragraphs[22])
				plantilla.paragraphs[19].replace('ha sido afectado por las interrupciones indicadas en la Tabla siguiente:','no ha sido afectado por interrupciones.')
		if self.parentTipo == PaseArchivo:
			if self.resolucionSup:
				if not self.accionesSumariales:
					self.borrarParrafo(plantilla.paragraphs[8])
					self.borrarParrafo(plantilla.paragraphs[8])
					self.borrarParrafo(plantilla.paragraphs[8])
					self.borrarParrafo(plantilla.paragraphs[8])
					self.borrarParrafo(plantilla.paragraphs[9])
					self.borrarParrafo(plantilla.paragraphs[9])
				else:
					self.borrarParrafo(plantilla.paragraphs[28])
					self.borrarParrafo(plantilla.paragraphs[28])
			if not self.auditadoSiNo:
				if not self.accionesSumariales:
					self.borrarParrafo(plantilla.paragraphs[8])
					self.borrarParrafo(plantilla.paragraphs[8])
				else:
					self.borrarParrafo(plantilla.paragraphs[11])
					self.borrarParrafo(plantilla.paragraphs[11])

				
		if self.cortes3Años:
			plantilla = self.completarTablaCortes3Años(plantilla,1)

		if self.parentTipo != CargosDemora:
			penalizaTipo['Producto'] = resultadosMediciones['P']
			penalizaTipo['Servicio'] = self.penalizaCortes
			
			inicio = 'Que en' if self.parentTipo == Resolucion else 'En'
			res = list(filter(lambda x: resultadosMediciones[x], resultadosMediciones))
			
			self.detallesTabla = '{} la tabla precedente, {}.'.format(inicio,self.resultadosMedicionesParrafo[res[0]]) if len(res)==1 else '{} la tabla precedente, {}. En igual sentido, {}.'.format(inicio,self.resultadosMedicionesParrafo[res[0]],self.resultadosMedicionesParrafo[res[1]]) if len(res)==2 else '{} la tabla precedente, {}. En igual sentido, {}. Además, {}.'.format(inicio,self.resultadosMedicionesParrafo[res[0]],self.resultadosMedicionesParrafo[res[1]],self.resultadosMedicionesParrafo[res[2]])
		
		if self.parentTipo not in  [PaseArchivo,CargosDemora]:
			self.motivoResolucion = self.motivoPenalizacion = 'la Calidad del ' + ' y '.join(filter(lambda x: penalizaTipo[x],penalizaTipo)) + ' Técnico'
		self.sumCortes = self.suministros[0]
		
		self.estacion = 'invernal {}'.format(localtime().tm_year) if localtime().tm_mon < 7 else 'estival {}/{}'.format(localtime().tm_year,localtime().tm_year+1)
		self.estacionSinAnio = self.estacion.split()[0]

		plantilla = self.replaceRuns(plantilla)
		
		self.saveDocx(plantilla,'{}/{} {} - IGG.docx'.format(self.dir,self.expedienteSup.replace('/','-'),self.usuarioSup))
		
		return self.__dict__
		
class ExcelCortes3A(Informe):
	"""
	Hace:
		planilla de cortes 3 años
	"""
	def __init__(self,parentDict):
		if parentDict:
			self.amChild = True
			self.__dict__ =  parentDict
		else:
			self.amChild = False
		global openpyxl
		import openpyxl
	
	def getSemestre(self,mes):
		"""
		devuelve el semestre a partir de "mes"
		"""
		sem = 1 if mes in [1,2,3,4,5,6] else 2 if mes in [7,8,9,10,11,12] else None
		return sem

	def setColor(self,columnas,fila):
		"""
		Sombrea las celdas producto de cada combinacion de los elementos en la LISTA "columnas" y "fila"
		"""
		for col in columnas:
			self.planilla['{}{}'.format(col,fila)].fill = openpyxl.styles.PatternFill(start_color='FF888888',end_color='FF888888',fill_type='solid')

	def setBorde(self,columnas,fila,bajo = False):
		"""
		Bordea las celdas producto de cada combinacion de los elementos en la LISTA "columnas" y "fila"
		Si "bajo" es True, el borde inferior de la celda se pone mas grueso
		"""
		bottom = self.thick if bajo else self.thin
		for col in columnas:
			self.planilla['{}{}'.format(col,fila)].border = openpyxl.styles.Border(top=self.thin, left=self.thin, right=self.thin, bottom=bottom)
			
	def getData(self):
		if not self.amChild:
			self.usuarioSup = input('Usuario:> ')
			print('Suministro:\nXXXXXXXXXXX\r',end='')
			self.sumCortes = input()
			self.getCortes3Años(self.sumCortes,tarifa = True)
		self.distribuidor = input('Distribuidor (prescindible):> ')
		self.ET = input('ET (prescindible):> ')
		
	def make(self):
		self.getData()
		if not self.amChild: self.mkdir(self.usuarioSup)
		print('Generando planilla de cortes 3 años...')	

		self.semestresPenalizados = []
		self.thin = openpyxl.styles.Side(border_style="thin", color="000000")
		self.thick = openpyxl.styles.Side(border_style="thick", color="000000")
		
		reCortesNoPenalizables = '(?i)FOT|FOD|FIU|FCL|CSC'
		header = 'Sum {} | Seta {} Dist. {} ET {} | {}'.format(self.sumCortes,self.seta,self.distribuidor,self.ET,self.usuarioSup)
		libroExcel = openpyxl.load_workbook(self.plantillas['Cortes3Años'])
		self.planilla = libroExcel['Expediente']
		
		self.planilla['A1'].value = header
		fila = 2
		cantidad = 0
		totalTiempo = 0
		mesInicio = strptime(self.fechascortes3Años[0],'%d/%m/%y').tm_mon
		año = strptime(self.fechascortes3Años[0],'%d/%m/%y').tm_year
		semestre = self.getSemestre(mesInicio)
		for corte in self.cortes3Años:
			self.planilla['A%d'%fila].value = self.sumCortes
			self.planilla['B%d'%fila].value = corte[1]
			self.planilla['C%d'%fila].value = corte[3].strftime('%d/%m/%Y %H:%M')
			self.planilla['D%d'%fila].value = corte[4].strftime('%d/%m/%Y %H:%M')
			self.setBorde(['A','B','C','D'],fila)
			duracionRaw = round((corte[4]-corte[3]).total_seconds()/60,10)
			duracion = round(duracionRaw)
			self.planilla['E%d'%fila].value = str(duracion)
			self.planilla['F%d'%fila].value = corte[6]
			self.setBorde(['E','F'],fila)

			semestreCorte = self.getSemestre(corte[3].month)
			añoCorte = corte[3].year
			if (semestreCorte != semestre) or (año != añoCorte):
				if totalTiempo > self.limites[self.tarifa]['duracion']/60.0 or cantidad > self.limites[self.tarifa]['cantidad']:
					self.setColor(['G','H','I'],(fila-1))
					self.semestresPenalizados.append('{}º Semestre/{}'.format(semestre,año))
				self.setBorde(['A','B','C','D','E','F','G','H','I'],(fila-1),bajo = True)
				self.planilla['G%d'%(fila-1)].value = str(cantidad)
				self.planilla['H%d'%(fila-1)].value = str(round(totalTiempo,1))
				self.planilla['I%d'%(fila-1)].value = '{}º sem {}'.format(semestre,año)
				semestre = semestreCorte
				año = añoCorte
				cantidad = 0
				totalTiempo = 0
			
			if duracionRaw > 3 and not search(reCortesNoPenalizables,corte[6]):
				self.setColor(['E'],fila)
				cantidad += 1
				totalTiempo += duracionRaw
			
			fila += 1
			
		if totalTiempo > self.limites[self.tarifa]['duracion']/60.0 or cantidad > self.limites[self.tarifa]['cantidad']:
			self.setColor(['G','H','I'],(fila-1))
			self.semestresPenalizados.append('{}º Semestre/{}'.format(semestre,año))
		self.planilla['G%d'%(fila-1)].value = str(cantidad)
		self.planilla['H%d'%(fila-1)].value = str(round(totalTiempo,1))
		self.planilla['I%d'%(fila-1)].value = '{}º sem {}'.format(semestre,año)
		self.setBorde(['A','B','C','D','E','F','G','H','I'],(fila-1),bajo = True)
		
		if not self.amChild:
			libroExcel.save('{}/{} {} - Cortes.xlsx'.format(self.dir,self.usuarioSup,self.sumCortes))
		else:
			libroExcel.save('{}/{} {} - Cortes.xlsx'.format(self.dir,self.expedienteSup.replace('/','-'),self.usuarioSup))
		
		return self.__dict__

class PrimeraNoPenalizada(Informe):			
	def make(self):
		igg = GerenciaGeneral(self.__dict__,type(self))
		self.__dict__ = igg.make()
		notaU = NotaUsuario(self.__dict__,type(self))
		self.__dict__ = notaU.make()
		notaD = NotaDistribuidora(self.__dict__,type(self))
		self.__dict__ = notaD.make()
		
		return self.__dict__

class NotaUsuario(Informe):
	def getData(self):
		print('Direccion Postal')
		self.direccionPostal = input(':> ')
		print('Departamento')
		self.departamentoSup = input(':> ')
	
	def make(self):
		print('Generando Nota al Usuario...')
		self.getData()
		if not self.amChild: self.mkdir()
		plantilla = self.openDocx(self.plantillas['NU'])
		plantilla = self.completarTablaMediciones(plantilla,0)
		if self.cortes:
			plantilla = self.completarTablaCortes(plantilla,1)
		else:
			if self.parentTipo == PrimeraNoPenalizada:
				self.borrarParrafo(plantilla.paragraphs[37])
				self.borrarTabla(plantilla.tables[1])
				self.borrarParrafo(plantilla.paragraphs[38])
				self.borrarParrafo(plantilla.paragraphs[38])
				plantilla.paragraphs[35].replace('ha sido afectado por las interrupciones indicadas en la Tabla siguiente:','no ha sido afectado por interrupciones.')
		self.estacion = 'invernal {}'.format(localtime().tm_year) if localtime().tm_mon < 7 else 'estival {}/{}'.format(localtime().tm_year,localtime().tm_year+1)
		self.estacionSinAnio = self.estacion.split()[0]
		plantilla = self.replaceRuns(plantilla)
		self.saveDocx(plantilla,'{}/{} {} - NU.docx'.format(self.dir,self.expedienteSup.replace('/','-'),self.usuarioSup))
				
		return self.__dict__
			
class NotaDistribuidora(Informe):
	def getData(self):
		print('Nota A.C.C. de referencia:')
		print('xxxx/xx\r',end='')
		self.notaReferencia = input()
	
	def make(self):
		print('Generando Nota la Distribuidora...')
		self.getData()
		if not self.amChild: self.mkdir()
		if self.parentTipo == PrimeraNoPenalizada:
			plantilla = self.openDocx(self.plantillas['ND1NP'])
		elif self.parentTipo == SegundaPenalizada:
			plantilla = self.openDocx(self.plantillas['ND2P'])
			self.tipoSuministrosPenalizados = 'el suministro' if len(self.suministrosPenalizados) == 1 else 'los suministros' if len(self.suministrosPenalizados) > 1 else None
			if self.tipoSuministrosPenalizados:
				if len(self.suministrosPenalizados) > 1:
					plural = True
					self.suministrosPenalizados = ', '.join(self.suministrosPenalizados[:-1]) + ' y ' + self.suministrosPenalizados[-1]
				elif len(self.suministrosPenalizados) == 1:
					plural = False
					self.suministrosPenalizados = self.suministrosPenalizados[0]
				med = 'las mediciones efectuadas' if plural else 'la medición efectuada'
				res = 'resultaron apartadas' if plural else 'resultó apartada'
				self.parrafoND = '{} en {} {}, {} de los límites establecidos en el Contrato de Concesión, en el periodo entre el {} y el {}'.format(med,self.tipoSuministrosPenalizados,self.suministrosPenalizados,res,self.fechaInicioMedicion,self.fechaFinMedicion)
				self.parrafoProducto = "Estudios de flujo de carga en baja tensión (firmados por el profesional responsable de dicha tarea), teniendo en cuenta los valores de máxima demanda registrados en el Distribuidor que alimenta la zona del reclamo, y los valores máximos proyectados para la próxima temporada {}, que acrediten la eficacia de las acciones de normalización a instrumentar, a partir de los cuales pueda concluirse que los niveles de calidad de producto técnico brindada en la zona del reclamo, serán adecuados a las previsiones requeridas en la normativa contractual de aplicación, (destacando que los datos de entrada y salida del referido estudio de flujo, deberán ser conservados y estar disponibles para su consulta).".format(self.estacion)

			else:
				self.parrafoND = 'la cantidad y/o duración de las interrupciones registradas en el referido suministro entre el {} y el {}, han superado los límites establecidos en el Contrato de Concesión'.format(self.fechasCortes[0],self.fechasCortes[1])
				self.borrarParrafo(plantilla.paragraphs[20])
				
			if self.penalizaCortes:
				self.parrafoServicio = "Informe de las acciones tendientes a mejorar {} (incluyendo informe de auditoría detallada de las condiciones de las redes de MT y BT a las que se vincula el Suministro mencionado en el reclamo, abarcando desde el punto de alimentación primario en la Estación Transformadora a la que se vincula el suministro, así como plan de ejecución de las acciones necesarias orientadas a mejorar la confiabilidad y calidad del servicio).".format(self.motivoPenalizacion)
			else:
				self.borrarParrafo(plantilla.paragraphs[18])
		self.estacion = 'invernal {}'.format(localtime().tm_year) if localtime().tm_mon < 7 else 'estival {}/{}'.format(localtime().tm_year,localtime().tm_year+1)
		self.estacionSinAnio = self.estacion.split()[0]
			
		plantilla = self.replaceRuns(plantilla)
		self.saveDocx(plantilla,'{}/{} {} - ND.docx'.format(self.dir,self.expedienteSup.replace('/','-'),self.usuarioSup))
				
		return self.__dict__

class SegundaPenalizada(Informe):
	def make(self):
		igg = GerenciaGeneral(self.__dict__,type(self))
		self.__dict__ = igg.make()
		notaD = NotaDistribuidora(self.__dict__,type(self))
		self.__dict__ = notaD.make()
		
		return self.__dict__

class PaseArchivo(Informe):
	def make(self):
		igg = GerenciaGeneral(self.__dict__,type(self))
		self.__dict__ = igg.make()

class CargosDemora(Informe):
	"""
	Hace nota a la distribuidora con formulacion de cargos
	e informe a gerencia general
	"""
	def make(self):
		igg = GerenciaGeneral(self.__dict__,type(self))
		self.__dict__ = igg.make()
		print('Generando Formulcion de Cargos...')
		plantilla = self.openDocx(self.plantillas['CargosDemora'])
		plantilla = self.completarTablaCortes(plantilla,0)
		plantilla = self.replaceRuns(plantilla)
		self.saveDocx(plantilla,'{}/{} {} - Formulacion de Cargos por Demora.docx'.format(self.dir,self.expedienteSup.replace('/','-'),self.usuarioSup))
		
		return self.__dict__
		
class Nota25m(Informe):
	def __init__(self,parentDict,parentTipo=None):
		self.plantilla = self.openDocx('Recursos/modelo25m.docx')
		self.dir = 'Notas 25m'
		self.amChild = False
		self.parentTipo = parentTipo
	
	def getData(self):
		notas = []
		opciones = ['solicitarle','reiterarle']
		
		self.iniciador = input('Iniciador\n:> ')
		print('Expediente:\n550.xxxx/xx',end='\r550.')
		self.expedienteSub = "550."+input()

		print('Notas A.C.C. #\nxxxx/xx',end='\r')
		while True:
			temp = input()
			if temp:
				notas.append(temp)
				print('xxxx/xx',end='\r')
			else: break
		if len(notas) > 1:
			self.plurSing = 'Notas'
			self.notasReferencia = ', '.join(notas[:-1])+' y '+notas[-1]
		else:
			self.plurSing = 'Nota'
			self.notasReferencia = notas[0]
			
		print('Resolución E.P.R.E. N°\nxxx/xx',end='\r')
		self.resolucionReferencia = input()
		self.tipoNota = opciones[int(input('1) Solicitar\n2) Reiterar\n:> '))-1]

		print('Notas de obra:')
		for i in range(len(notas)):
			print('{}) {}'.format(i+1,notas[i]))
		notaDeObra = [x.strip().lstrip() for x in input(':> ').split(',')]
		if len(notaDeObra) == 1:
			if '' in notaDeObra[0] or notaDeObra[0] == 1:
				self.obraPlSing = 'la Nota A.C.C. #'
				self.notaDeObrasFormato = notas[0]
		else:
			self.obraPlSing = 'las Notas A.C.C. #'
			self.notaDeObrasFormato = ', '.join(notas[:-1])+' y '+notas[-1]

		self.observaciones = input('Detalles de obras\n:> ')
		if self.observaciones:
			self.observaciones = ' (' + self.observaciones + ')'
			
		print('Plazo:> 5 dias',end='\rPlazo:> ')
		self.plazoNumero = input()
		if not self.plazoNumero:
			self.plazoNumero = '5'
		self.plazoPalabra = plazos[self.plazoNumero]
	
	def make(self):
		self.getData()
		if not self.resolucionReferencia:
			self.borrarParrafo(self.plantilla.paragraphs[4])
		self.plantilla = self.replaceRuns(self.plantilla)
		self.saveDocx(self.plantilla,'{}/{} {} Nota 25 inciso m y contancias de culminación de obra.docx'.format(self.dir,self.expedienteSub.replace('/','-'),self.iniciador))
		
class NotaCPP(Informe):
	def __init__(self,parentDict,parentTipo=None):
		self.plantilla = self.openDocx('Recursos/NotaCPP.docx')
		self.dir = 'Notas Respuesta a CPP'
		self.amChild = False
		self.parentTipo = parentTipo	
	
	def getData(self):
		print('Expediente:\n550.xxxx/xx',end='\r550.')
		self.expedienteSub = '550.'+input()
		
		notas = []
		print('Notas A.C.C. #\nxxxx/xx',end='\r')
		while True:
			temp = input()
			if temp:
				notas.append(temp)
				print('xxxx/xx',end='\r')
			else: break
		if len(notas) > 1:
			self.plurSing = 'Notas'
			self.notasReferencia = ', '.join(notas[:-1])+' y '+notas[-1]
		else:
			self.plurSing = 'Nota'
			self.notasReferencia = notas[0]
		
		if len(notas) > 1:
			print('Nota Respuesta CPP:')
			for index,nota in enumerate(notas,start=1):
				print('{}) {}'.format(index,nota))
			index = int(input(':> '))
			self.notaConcurso = notas[index-1]
		else: self.notaConcurso = notas[0]
			
		print('Resolución E.P.R.E. N°\nxxx/xx',end='\r')
		self.resolucionReferencia = input()
			
		print('Plazo:> 5 dias',end='\rPlazo:> ')
		self.plazoNumero = input()
		if not self.plazoNumero:
			self.plazoNumero = '5'
		self.plazoPalabra = plazos[self.plazoNumero]
	
	def make(self):
		self.getData()
		if not self.resolucionReferencia:
			self.borrarParrafo(self.plantilla.paragraphs[5])
		self.plantilla = self.replaceRuns(self.plantilla)
		self.saveDocx(self.plantilla,'{}/Respuesta a CPP - {}.docx'.format(self.dir,self.expedienteSub.replace('/','-')))

class AvanceDeObra(Informe):
	def __init__(self,parentDict,parentTipo=None):
		self.plantilla = self.openDocx('Recursos/avanceDeObra.docx')
		self.dir = 'Notas Avance de Obra'
		self.amChild = False
		self.parentTipo = parentTipo

	def getData(self):
		self.iniciador = input('Iniciador\n:> ')
		print('Expediente:\n550.xxxx/xx',end='\r550.')
		self.expedienteSub = '550.'+input()

		print('Nota A.C.C. #\nxxxx/xx',end='\r')
		self.notaReferencia = input()
			
		print('Resolución E.P.R.E. N°\nxxx/xx',end='\r')
		self.resolucionReferencia = input()
			
		print('Plazo:> 5 dias',end='\rPlazo:> ')
		self.plazoNumero = input()
		if not self.plazoNumero:
			self.plazoNumero = '5'
		self.plazoPalabra = plazos[self.plazoNumero]
	
	def make(self):
		self.getData()
		if not self.resolucionReferencia:
			self.borrarParrafo(self.plantilla.paragraphs[5])
		self.plantilla = self.replaceRuns(self.plantilla)
		self.saveDocx(self.plantilla,'{}/{} {} Avance de obra.docx'.format(self.dir,self.expedienteSub.replace('/','-'),self.iniciador))
		
class DebidaRespuesta(Informe):
	def __init__(self,parentDict,parentTipo=None):
		self.plantilla = self.openDocx('Recursos/DebidaRespuesta.docx')
		self.dir = 'Notas D. Respuesta'
		self.amChild = False
		self.parentTipo = parentTipo
	
	def getData(self):
		notas = []
		notasEpre = []
		
		self.iniciador = input('Iniciador\n:> ')
		print('Expediente:\n550.xxxx/xx',end='\r550.')
		self.expedienteSub = "550."+input()

		print('Notas E.P.R.E. #\nxxxx/xx',end='\r')
		while True:
			temp = input()
			if temp:
				notasEpre.append(temp)
				print('xxxx/xx',end='\r')
			else: break
		if len(notasEpre) > 1:
			self.plurSingEpre = 'Notas'
			self.notasEpre = ', '.join(notasEpre[:-1])+' y '+notasEpre[-1]
		else:
			self.plurSingEpre = 'Nota'
			self.notasEpre = notasEpre[0]
		
		print('Notas A.C.C. #\nxxxx/xx',end='\r')
		while True:
			temp = input()
			if temp:
				notas.append(temp)
				print('xxxx/xx',end='\r')
			else: break
		if len(notas) > 1:
			self.plurSing = 'Notas'
			self.notasReferencia = ', '.join(notas[:-1])+' y '+notas[-1]
		else:
			self.plurSing = 'Nota'
			self.notasReferencia = notas[0]
			
		print('Resolución E.P.R.E. N°\nxxx/xx',end='\r')
		self.resolucionReferencia = input()

		print('Notas a responder:')
		for i in range(len(notasEpre)):
			print('{}) {}'.format(i+1,notasEpre[i]))
		notaEPRE = [x.strip().lstrip() for x in input(':> ').split(',')]
		if len(notaEPRE) == 1:
			if '' in notaEPRE[0] or notaEPRE[0] == 1:
				self.respPlSing = 'la Nota'
				self.notaRespFormato = notasEpre[0]
		else:
			self.respPlSing = 'las Notas'
			self.notaRespFormato = ', '.join(notasEpre[:-1])+' y '+notasEpre[-1]

		print('Plazo:> 5 dias',end='\rPlazo:> ')
		self.plazoNumero = input()
		if not self.plazoNumero:
			self.plazoNumero = '5'
		self.plazoPalabra = plazos[self.plazoNumero]
	
	def make(self):
		self.getData()
		if not self.resolucionReferencia:
			self.borrarParrafo(self.plantilla.paragraphs[3])
		self.plantilla = self.replaceRuns(self.plantilla)
		self.saveDocx(self.plantilla,'{}/{} {} Nota Debida Respuesta.docx'.format(self.dir,self.expedienteSub.replace('/','-'),self.iniciador))

@manejarErrores
def main():
	listaDisponible = {	1:Resolucion,
				2:ExcelCortes3A,
				3:PrimeraNoPenalizada,
				4:SegundaPenalizada,
				5:CargosDemora,
				6:PaseArchivo,
				7:Nota25m,
				8:NotaCPP,
				9:AvanceDeObra,
				10:DebidaRespuesta}
	print('Tipo de informe:\n1) Resolucion\n2) Planilla Excel 3 Años\n3) Primera Medicion No Penalizada\n4) Segunda Medicion Penalizada\n5) Formulacion de Cargos por Demora\n6) Pase a Archivo\n7) Nota 25m\n8) Nota CPP\n9) Avance de Obra\n10) Debida Respuesta')
	inf = listaDisponible[int(input(':> '))](None)
	try: inf.make()
	except KeyboardInterrupt: exit()
	input('\nListo')

if __name__ == '__main__':
	main()