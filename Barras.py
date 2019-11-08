import pyodbc
import openpyxl
from os import walk, getcwd, path, mkdir, remove
from re import search,sub
from sys import argv
from docx import Document
from time import strptime, strftime, sleep, mktime, time, ctime, asctime, localtime
from rutas import unitariosDoc,plantillaSemestral,plantillaCortesSemestral,plantillaBarras
from shutil import copyfile
from calendar import monthrange
from statistics import mean
from Informacion import valoresLimite,meses,mesesRev,mesesR32,sheets,factoresDeInversion,choices,posicionesGraficosBarras,manejarErrores
from docx.shared import Pt
from openpyxl.chart import LineChart,Reference
from openpyxl.chart.text import RichText
from openpyxl.drawing.text import Paragraph, ParagraphProperties, CharacterProperties, Font

filaErrSemestral = 3
filaDatSemestral = 16
colores = {1:'FFC0C0C0',} # no es una gallina, es el código para el GRIS MEDIO

class ErrorDat(Exception): pass
class ErrorNombreBarra(Exception):pass
class ErrorMesBarra(Exception):pass

class Corte():
	def __init__(self,linea):
		groups = search('(?s)(?P<fecha>^\d\d/\d\d/\d{2,4}).*(?P<hora>(?<=\s)\d\d:\d\d:\d\d)(?=\s|\.)(?:\s|.0{3}\s)(?P<detalle>.+)',linea.strip())
		self.fecha = groups.group('fecha')
		self.hora = groups.group('hora')
		try: self.fechaTupla = strptime(' '.join([self.fecha,self.hora]),'%d/%m/%y %H:%M:%S')
		except: self.fechaTupla = strptime(' '.join([self.fecha,self.hora]),'%d/%m/%Y %H:%M:%S')
		self.detalle = groups.group('detalle')
		#False si se corte, True si vuelve
		self.tipo = True if search('(?i)Vuelta\sde|Fin\sINT',self.detalle) else False
		
class Registro():
	def __init__(self,reg,dat):
		self.penaliza = False
		self.fecha = reg[0]
		self.horario = reg[1]
		try: self.fechaTupla = strptime(' '.join([self.fecha,self.horario]),'%d/%m/%y %H:%M')
		except:
			try: self.fechaTupla = strptime(' '.join([self.fecha,self.horario]),'%d/%m/%y %H:%M:%S')
			except:
				try: self.fechaTupla = strptime(' '.join([self.fecha,self.horario]),'%d/%m/%Y %H:%M')
				except: self.fechaTupla = strptime(' '.join([self.fecha,self.horario]),'%d/%m/%Y %H:%M:%S')
		self.hora = int(reg[1].split(':')[0])
		self.V1 = float(reg[2])
		self.V1max = float(reg[3])
		self.V1min = float(reg[4])
		try:
			self.V2 = float(reg[8])
			self.V2max = float(reg[9])
			self.V2min = float(reg[10])
			self.V3 = float(reg[14])
			self.V3max = float(reg[15])
			self.V3min = float(reg[16])
			self.thd = float(reg[20])
			self.flicker = float(reg[21])
			self.energia = float(reg[23])
			self.Vmax = max([self.V1max,self.V2max,self.V3max])
			self.Vmin = max([self.V1min,self.V2min,self.V3min])
			try: self.anormalidad = True if reg[24] == 'A' else None
			except: self.anormalidad = False
			self.corrupto = False
		except (IndexError,ValueError):
			print('EL dat {} tiene menos columnas que las necesarias'.format(dat))
			if __name__ == '__main__':salir()
			else: self.corrupto = True
		self.desequilibrio = None

class Dat():
	def __init__(self,pathToDat,datName,mes,barra):
		self.name = datName
		self.intrusos = []
		self.registros = []
		self.timestamps = []
		self.mesDat = mes.mes
		try:
			try:
				with open(pathToDat+datName,'r') as datRaw:
					dat = datRaw.readlines()
			except UnicodeDecodeError:
				with open(pathToDat+datName,'r',encoding='latin-1') as datRaw:
					dat = datRaw.readlines()

			paraIntervalo = sub('\s+',' ',dat[1]).split()
			for id,i in enumerate(paraIntervalo):
				if 'intervalo' in i.lower() or '::' in i:
					periodo,unidad = paraIntervalo[id+1],paraIntervalo[id+2]
					break
			periodo = int(periodo) if unidad == 'min.' else int(periodo)*60 if unidad == 'seg.' else None

			try:
				self.tv = float(search('(?<=Factor de Corrección: )\d+,\d+',dat[2].split('\t')[3].strip().lstrip()).group().replace(',','.'))
				self.ti = float(search('(?<=Factor de Corrección: )\d+,\d+',dat[3].split('\t')[3].strip().lstrip()).group().replace(',','.'))
				if self.tv != barra.tvESJ or self.ti != barra.tiESJ:
					print('El archivo {} se procesó con TV: {} y TI: {}. Valores correctos: TV: {} y TI: {}'.format(self.name,self.tv,self.ti,barra.tvESJ,barra.tiESJ))
					raise ErrorDat
			except ErrorDat: raise ErrorDat
			except:
				print('No se encontró el factor de tensión o corriente en el dat',self.name)
				raise ErrorDat
			
			for line in dat[9:]:
				registro = [y.strip().replace(',','.') for y in line.strip('\n').split('\t')]
				reg = Registro(registro,datName)
				if reg.fechaTupla.tm_mon != mes.mesInt:
					self.intrusos.append(reg)
					continue
				elif reg.corrupto: raise ErrorDat
				else:
					self.registros.append(reg)
					self.timestamps.append(mktime(reg.fechaTupla))
			
			tmp = self.registros[::]
			if all([x.anormalidad for x in self.registros]):
				for reg in tmp: reg.anormalidad = False
				self.registros = tmp
			else:		
				for reg in tmp:
					if reg.anormalidad: self.registros.remove(reg)
				# anormalidadPrevia = False
				# anormalidadPosterior = False
				# primero = True
						
				# tmp = self.registros[::]
				# for reg in tmp:
					# if reg.anormalidad and primero:
						# try: anormalidadPosterior = self.registros[self.registros.index(reg)+1].anormalidad
						# except: pass
						# if self.registros.index(reg) != 0: anormalidadPrevia = self.registros[self.registros.index(reg)-1].anormalidad

						# if anormalidadPosterior: self.registros.remove(reg)
						# else:
							# self.registros.remove(reg)
							# primero = False
							
				# anormalidadPrevia = False
				# anormalidadPosterior = False
				# primero = True
				# self.registros.reverse()
				# tmp = self.registros[::]
				# for reg in tmp:
					# if reg.anormalidad and primero:				
						# try: anormalidadPosterior = self.registros[self.registros.index(reg)+1].anormalidad
						# except: pass
						# if self.registros.index(reg) != 0: anormalidadPrevia = self.registros[self.registros.index(reg)-1].anormalidad

						# if anormalidadPosterior: self.registros.remove(reg)
						# else: primero = False
				# self.registros.reverse()
			
				# tmp = self.registros[::]
				# for reg in tmp:
					# if not all([reg.V1,reg.V2,reg.V3,reg.V1max,reg.V2max,reg.V3max,reg.V1min,reg.V2min,reg.V3min,]) and reg.anormalidad:
						# self.registros.remove(reg)
			
			self.fases 			= 3
			self.periodo 		= periodo
			self.unidad			= unidad
			self.problemas = False
			self.setTimeStamps(barra)
			
		except FileNotFoundError:
			print('No existe el archivo {} para la barra {} en el mes {}'.format(datName,[barra.id,barra.puntoDeMedicion,barra.nombre],mes.mes))
			self.problemas = True
		except ErrorDat:
			self.problemas = True

	def setTimeStamps(self,barra,sincronizar=False):
		try:
			try:
				fechaInicio = strftime('%d/%m/%Y',self.registros[0].fechaTupla)
				fechaFin = strftime('%d/%m/%Y',self.registros[-1].fechaTupla)
			except IndexError:
				print('El archivo {} no tiene registros para el mes de {}'.format(self.name,self.mesDat,[barra.id,barra.puntoDeMedicion,barra.nombre,self.name]))
				if sincronizar: raise ErrorDat
				else: return
			horaInicio = strftime('%I:%M:%S %p',self.registros[0].fechaTupla)
			horaFin = strftime('%I:%M:%S %p',self.registros[-1].fechaTupla)
			self.horaInicio 	= horaInicio
			self.fechaInicio 	= fechaInicio
			self.horaFin 		= horaFin
			self.fechaFin		= fechaFin
		except ErrorDat: self.problemas = True
		
class Err():
	def __init__(self,pathToErr,errName,mes,barra):
		self.name = errName
		self.registros = []
		self.intrusos = []
		try:
			with open(pathToErr+errName,'r') as errRaw:
				err = errRaw.readlines()
				for line in err:
					if 'Inicio DIP' in line:
						self.serie1612 = True
						break
				else: self.serie1612 = False

				for line in err:
					corte = Corte(line)
					corte.errName = self.name
					corte.serie1612 = self.serie1612
					if corte.fechaTupla.tm_mon != mes.mesInt: self.intrusos.append(corte)
					else: self.registros.append(corte)
			self.problemas = False
		except FileNotFoundError:
			print('No existe el archivo {} para la barra {} en el mes {}'.format(errName,[barra.id,barra.puntoDeMedicion,barra.nombre],mes.mes))
			self.problemas = True

class Mes():
	def __init__(self,mes):
		self.mes = mes.lower().capitalize()
		self.mesInt = int(meses[mes.lower()])
		self.mesStr = str(self.mesInt).zfill(2)
		self.barras = []

class Barra():
	def __init__(self,mes,fila):
		# Mes
		self.mes				= mes
		# Datos Barra
		self.id 				= fila[0].value
		self.idNumero			= search('\d+',self.id).group()
		self.puntoDeMedicion 	= int(fila[1].value)
		self.nombre 			= fila[2].value
		self.direccion 			= fila[3].value
		self.localidad 			= fila[4].value
		self.tensionNominal 	= float(fila[5].value)
		self.tvESJ				= float(fila[6].value)
		self.tiESJ 				= float(fila[7].value)
		self.archivo 			= fila[8].value.split('.')[0] if fila[8].value else ''
		if not search('M[\d,N,D,O]+',self.archivo):
			print('La barra {} no tiene archivo para el mes de {}'.format([self.id,self.puntoDeMedicion,self.nombre],mes))
			self.archivo = None
		elif getMes(self.archivo)[0] != self.mes.capitalize():
			salir('El archivo {}.dat tiene nombre incorrecto.\nCorregirlo en el listado que envió ESJ para el mes de {}'.format(self.archivo,mes))
			self.archivo = None
		else:
			self.r32 = self.archivo+'.R32'
		self.resultadoInformado = fila[9].value
		self.observaciones 		= fila[10].value
		# Limites
		self.lInf			= self.tensionNominal*valoresLimite['qBarras']['baja']
		self.lSup			= self.tensionNominal*valoresLimite['qBarras']['alta']
		self.lSupPen 		= self.tensionNominal*(1+valoresLimite['penBarras'])
		self.lInfPen 		= self.tensionNominal*(1-valoresLimite['penBarras'])
		self.lFlicker 		= valoresLimite['flicker']
		self.lThd 			= valoresLimite['thdBarras']
		self.maximoRegistrosFueraDeRango = valoresLimite['Maximo porcentaje de registros fuera de rango']
		# Misc
		self.preguntar = False
		# Cortes
		self.cortes = []

	def procesarDat(self):
		registrosCortes = []
		registrosCorrectos,registrosPenalizados = [],[]
		registrosSubTension,registrosSobreTension = 0,0
		registrosSubTensionF1,registrosSubTensionF2,registrosSubTensionF3 = 0,0,0
		registrosSobreTensionF1,registrosSobreTensionF2,registrosSobreTensionF3 = 0,0,0
		energiasPenalizadasSobre,energiasPenalizadasSub = 0,0
		registrosSubTensionCalc,registrosSobreTensionCalc = [],[]
		flickerFueraDeRango,thdFueraDeRango = [],[]
		desequilibrios = []
		
		self.dat.promedioVF1,self.dat.promedioVF2,self.dat.promedioVF3 = map(mean,[[reg.V1 for reg in self.dat.registros],[reg.V2 for reg in self.dat.registros],[reg.V3 for reg in self.dat.registros]])
		self.dat.promedioVtotal = mean([self.dat.promedioVF1,self.dat.promedioVF2,self.dat.promedioVF3])
		self.dat.energia = 0.0
		
		tmp = self.dat.registros[::]
		for reg in self.dat.registros:
			minimo,maximo = min(reg.V1,reg.V2,reg.V3),max(reg.V1,reg.V2,reg.V3)
			if minimo < self.lInf or maximo > self.lSup:
				registrosCortes.append(reg)
				tmp.remove(reg)
		
		self.dat.registros = tmp[::]
		for reg in self.dat.registros:
			minimo,maximo = min(reg.V1,reg.V2,reg.V3),max(reg.V1,reg.V2,reg.V3)
			self.dat.energia += reg.energia
			
			registrosSubTensionF1 = registrosSubTensionF1+1 if reg.V1 < self.lInfPen else registrosSubTensionF1
			registrosSubTensionF2 = registrosSubTensionF2+1 if reg.V2 < self.lInfPen else registrosSubTensionF2
			registrosSubTensionF3 = registrosSubTensionF3+1 if reg.V3 < self.lInfPen else registrosSubTensionF3
			registrosSobreTensionF1 = registrosSobreTensionF1+1 if reg.V1 > self.lSupPen else registrosSobreTensionF1
			registrosSobreTensionF2 = registrosSobreTensionF2+1 if reg.V2 > self.lSupPen else registrosSobreTensionF2
			registrosSobreTensionF3 = registrosSobreTensionF3+1 if reg.V3 > self.lSupPen else registrosSobreTensionF3

			if minimo < self.lInfPen:
				registrosSubTension += 1
				energiasPenalizadasSub += reg.energia
				reg.penaliza = True
			if maximo > self.lSupPen:
				registrosSobreTension += 1
				energiasPenalizadasSobre += reg.energia
				reg.penaliza = True
			if not reg.penaliza: registrosCorrectos.append(reg)
			else: registrosPenalizados.append(reg)
				
			d1,d2,d3 = map(lambda x: abs(x-self.tensionNominal)/self.tensionNominal,[reg.V1,reg.V2,reg.V3])
			desequilibrios.append([d1,d2,d3])
			
			if reg.flicker > self.lFlicker:
				flickerFueraDeRango.append(reg.flicker)
			if reg.thd > self.lThd:
				thdFueraDeRango.append(reg.thd)
		
		registros = registrosPenalizados+registrosCorrectos
		totalRegistros = len(registros)
		totalRegistrosPenalizados = registrosSobreTension+registrosSubTension
		
		# calculo de multa	
		self.calcularMulta(totalRegistros,registrosPenalizados)

		self.dat.penaliza = None
		if 'fallida' not in self.__dict__:
			try:
				if totalRegistrosPenalizados/totalRegistros > self.maximoRegistrosFueraDeRango:
					self.dat.penaliza = True
				else:
					self.dat.penaliza = False
					self.dat.multaDown = 0
					self.dat.multaUp = 0
					self.dat.multaTotal = 0
				self.dat.fallida = None
			except ZeroDivisionError:
				print('El dat {} es posible fue procesado como BT. Reprocesar el archivo R32'.format(self.dat))
				self.dat.fallida = True
		
		self.dat.penalizaFlicker = True if len(flickerFueraDeRango) > totalRegistros*0.05 else False
		self.dat.penalizaThd = True if len(thdFueraDeRango) > totalRegistros*0.05 else False
			
		DF1,DF2,DF3 = [x[0] for x in desequilibrios],[x[1] for x in desequilibrios],[x[2] for x in desequilibrios]
		promediosDesequilibrios = list(map(mean,[DF1,DF2,DF3]))
		try: desequilibriosMaximos = list(map(max,[DF1,DF2,DF3]))
		except ValueError: desequilibriosMaximos = [0,0,0]
		
		self.dat.apartamientoMaximo 								= max(desequilibriosMaximos)
		self.dat.apartamientoMaximoF1 								= max(desequilibriosMaximos)
		self.dat.apartamientoPromedio 								= max(promediosDesequilibrios)
		self.dat.apartamientoPromedioF1 							= promediosDesequilibrios[0]
		self.dat.energiaPenalizada 									= self.dat.energiaPenalizadaUp+self.dat.energiaPenalizadaDown
		self.dat.energiaSobretension 								= self.dat.energiaPenalizadaUp
		self.dat.energiaSubtension 									= self.dat.energiaPenalizadaDown
		self.dat.energiaTotal 										= self.dat.energia
		self.dat.energiaTotalF1 									= self.dat.energia
		self.dat.flicker 											= mean([x.flicker for x in self.dat.registros])
		self.dat.flickerF1 											= self.dat.flicker
		self.dat.flickerFueraDeRango 								= len(flickerFueraDeRango)
		self.dat.flickerFueraDeRangoF1 								= len(flickerFueraDeRango)
		self.dat.flickerFueraDeRangoF2, self.dat.flickerFueraDeRangoF3 	= '0,000','0,000'
		self.dat.flickerPenalizable 								= 'SI' if self.dat.penalizaFlicker else 'NO'
		self.dat.flickerPenalizableF1 								= 'SI' if self.dat.penalizaFlicker else 'NO'
		self.dat.multaFueraDeRango 									= self.dat.multaTotal
		self.dat.promedioTension 									= self.dat.promedioVtotal
		self.dat.promedioTensionF1 									= self.dat.promedioVF1
		self.dat.promediosDesequilibrios 							= promediosDesequilibrios
		self.dat.resultado											= 'Penalizada' if self.dat.penaliza else 'No Penalizada' if not self.dat.fallida else 'Fallida'
		self.dat.tensionMaxima 										= max([x.Vmax for x in self.dat.registros])
		self.dat.tensionMaximaF1 									= max([x.V1max for x in self.dat.registros])
		self.dat.tensionMinima 										= min([x.Vmin for x in self.dat.registros])
		self.dat.tensionMinimaF1 									= min([x.V1min for x in self.dat.registros])
		self.dat.thdF1 												= mean([x.thd for x in self.dat.registros])
		self.dat.thdF2 												= '-'
		self.dat.thdF3 												= '-'
		self.dat.thdFueraDeRango 									= len(thdFueraDeRango)
		self.dat.thdFueraDeRangoF1 									= len(thdFueraDeRango)
		self.dat.thdPenalizable 									= 'SI' if self.dat.penalizaThd else 'NO'
		self.dat.thdPenalizableF1 									= self.dat.thdPenalizable
		self.dat.thdTotal 											= self.dat.thdF1
		self.dat.tipoMedicion 										= 'BARRAS'
		self.dat.tipoMulta 											= 'Ambos' if (self.dat.multaUp and self.dat.multaDown) else 'Sub' if self.dat.multaDown else 'Sobre' if self.dat.multaUp else ''
		self.dat.totalRegistros 									= totalRegistros
		self.dat.totalRegistrosF1 									= totalRegistros
		self.dat.totalRegistrosPenalizados 							= totalRegistrosPenalizados
		self.dat.totalRegistrosPenalizadosF1 						= registrosSobreTensionF1+registrosSubTensionF1
		self.dat.totalRegistrosSobretension 						= registrosSobreTension
		self.dat.totalRegistrosSobretensionF1 						= registrosSobreTensionF1
		self.dat.totalRegistrosSubtension 							= registrosSubTension
		self.dat.totalRegistrosSubtensionF1 						= registrosSubTensionF1
		self.dat.flickerPenalizableF2,self.dat.flickerPenalizableF3 	= 'NO','NO'
		self.dat.promedioTensionF2 									= self.dat.promedioVF2
		self.dat.promedioTensionF3									= self.dat.promedioVF3
		self.dat.tensionMaximaF2 									= max([x.V2max for x in self.dat.registros])
		self.dat.tensionMaximaF3 									= max([x.V3max for x in self.dat.registros])
		self.dat.tensionMinimaF2 									= min([x.V2min for x in self.dat.registros])
		self.dat.tensionMinimaF3 									= min([x.V3min for x in self.dat.registros])
		self.dat.thdPenalizableF2,self.dat.thdPenalizableF3 			= 'NO','NO'
		self.dat.thdFueraDeRangoF2,self.dat.thdFueraDeRangoF3 			= '0,000','0,000'
		self.dat.totalRegistrosF2 									= totalRegistros
		self.dat.totalRegistrosF3									= totalRegistros
		self.dat.apartamientoMaximoF2 								= desequilibriosMaximos[1]
		self.dat.apartamientoMaximoF3 								= desequilibriosMaximos[2]
		self.dat.apartamientoPromedioF2 							= promediosDesequilibrios[1]
		self.dat.apartamientoPromedioF3 							= promediosDesequilibrios[2]
		self.dat.totalRegistrosPenalizadosF2 						= registrosSobreTensionF2+registrosSubTensionF2
		self.dat.totalRegistrosPenalizadosF3 						= registrosSobreTensionF3+registrosSubTensionF3
		self.dat.totalRegistrosSobretensionF2 						= registrosSobreTensionF2
		self.dat.totalRegistrosSobretensionF3 						= registrosSobreTensionF3
		self.dat.totalRegistrosSubtensionF2 						= registrosSubTensionF2
		self.dat.totalRegistrosSubtensionF3 						= registrosSubTensionF3

	def procesarErr(self):
		cortes1612 = {'1':{'inicio':{},'fin':{}},'2':{'inicio':{},'fin':{}},'3':{'inicio':{},'fin':{}}}
		indexInicio = {'1':0,'2':0,'3':0}
		inicio,fin = False,False
		def duracion(inicio,fin):
			inicio,fin = mktime(inicio),mktime(fin)
			duracion = 0.0
			duracion = int(fin-inicio)
			return duracion
			
		cortes = list(filter(lambda x: True if search('(?i)Corte de Tensión|Vuelta de Tensión|(?:(?<=U)\d(?=\s)).*(?:Fin\sINTERRRUPT|Inicio\sINTERRUPT)(?:\s\d\d:\d\d:\d\d)?',x.detalle) else False,self.err.registros))
		cortes.sort(key = lambda x:mktime(x.fechaTupla))
		
		i = 0
		for corte in cortes:
			if self.err.serie1612:
				detalle = search('(?i)(?P<fase>(?<=U)\d(?=\s)).*(?P<det>Fin\sINTERRRUPT|Inicio\sINTERRUPT)(?P<duracion>\s\d\d:\d\d:\d\d)?',corte.detalle)
				if corte.tipo:
					corte.faseFin = detalle.group('fase')
					corte.detalle = ' '.join((corte.faseFin,detalle.group('det')))
					corte.duracion = detalle.group('duracion').strip()
					cortes1612[corte.faseFin]['fin'] = corte
				else:
					corte.faseInicio = detalle.group('fase')
					corte.detalle = ' '.join(('Fase',corte.faseInicio,detalle.group('det')))
					cortes1612[corte.faseInicio]['inicio'] = corte

				for f in range(1,4):
					if cortes1612[str(f)]['inicio'] and cortes1612[str(f)]['fin']:
						st = cortes1612[str(f)]['inicio']
						nd = cortes1612[str(f)]['fin']
						st.duracion = nd.duracion
						self.cortes.append((st,nd))
						cortes1612[str(f)] = {'inicio':{},'fin':{}}
			else:
				durMinutos = 0
				if corte.tipo:
					fin = True
					nd = corte
				else:
					inicio = True
					st = corte
				if inicio and fin:
					nd.duracion = duracion(st.fechaTupla,nd.fechaTupla)
					st.duracion = nd.duracion
					self.cortes.append((st,nd))
					inicio,fin = None,None
					st,nd = None,None
			i+=1

	def calcularMulta(self,totalRegistros,registrosPenalizados):
		multaTotal =			0
		multaUp = 				0
		multaDown = 			0
		energiaPenalizadaUp = 	0
		energiaPenalizadaDown = 0
		for reg in registrosPenalizados:
			unitarios = valoresUnitariosPost if mktime(reg.fechaTupla) > mktime(fechaCambioUnitarios) else valoresUnitariosPrevios
		
			v,desvioPorcentual = max(map(lambda x: (x,abs(x-self.tensionNominal)/self.tensionNominal),[reg.V1,reg.V2,reg.V3]),key= lambda x:x[1])# Desvios porcentuales
			difTensiones = v-self.tensionNominal
			
			desviaciones = sorted(unitarios.keys())+[1]
			for desviacion in desviaciones:
				if desvioPorcentual > desviacion:
					precioKilowatt = unitarios[desviacion]
				else: break
			else: precioKilowatt = 0.0
			
			multaParcial = reg.energia*precioKilowatt
			multaTotal += multaParcial
			if difTensiones > 0:
				multaUp += multaParcial
				energiaPenalizadaUp += reg.energia
			else:
				multaDown += multaParcial
				energiaPenalizadaDown += reg.energia

		self.dat.multaTotal = multaTotal
		self.dat.multaUp = multaUp
		self.dat.multaDown = multaDown
		self.dat.energiaPenalizadaUp = energiaPenalizadaUp
		self.dat.energiaPenalizadaDown = energiaPenalizadaDown

def salir(mensaje=''):
	"""
	Imprime el mensaje que se le entregue como argumento y termina el programa.
	"""
	print(mensaje,'\nEnter para terminar...')
	input()
	exit()

def fillColor(code):
	fill = openpyxl.styles.PatternFill(start_color=colores[code],
                   end_color=colores[code],
                   fill_type='solid')
	return fill

def seekExpansivo(iterable,startIndex):
	puntero = 0
	while puntero <= (len(iterable)+startIndex):
		try: yield iterable[startIndex+puntero]
		except IndexError: pass
		try:
			b = startIndex-puntero
			if b>=0 and b!=startIndex: yield iterable[b]
		except IndexError: pass
		puntero+=1

def cellFontStyle(style):
	if style == 'bold':
		style = openpyxl.styles.Font(bold=True)
	return style
	
def getDir():
	"""
	Define un global con el Directorio donde se va a trabajar.
	No devuelve nada porque "dir" es global.
	Si no se le pone la direccion, el programa pone como default el directorio
	donde se encuentra a la hora de ejecución.
	"""
	global dir
	check1,check2 = False,False
	anotacion = ''
	dir = input('Carpeta:> ')
	print()
	if not dir:	dir = getcwd()
	else:
		for root,dirs,files in walk(dir):
			if set([dir.lower() for dir in dirs]).intersection(set(meses.keys())): check1 = True
			if all([not x.endswith('.dat') for x in files]): check2 = True
		if not check1:
			print('El directorio que ingresaste no contiene meses a procesar.')
		if not check2:
			print('Es posible que hayas ingresado el directorio de algún mes en especial.')
			anotacion = '<<< Posiblemente hayas ingresado este.'
		if not check1 or not check2:
			print("""
X Semestre 20XX:  <<< Este es el directorio que hay que ingresar
	|__> Mes1
		|__> Archivos Enviados {}
			|__> M0119560.R32
			|__> M0119560.dat
			|__> M0119560.err
			|__> ...
	|__> Mes...
	|__> Semestral
			""".format(anotacion))
			salir()

def getBarraByFile(filename):
	barra = filename[1:3]
	if barra.isnumeric(): return filename[1:3]
	else: raise ErrorNombreBarra
		
def getFiles(dir,*args):
	"""
	devuelve una lista de listas con los archivos en el directorio de argumento
	"""
	archivos = []
	for root, folder, file in walk('/'.join([dir]+list(args))):
		archivos.append(file)
	return archivos

def getMes(file):
	"""
	devuelve el mes al que pertenece el archivo de argumento
	Ej: M01D81715.dat --> ('Diciembre',12,'D')
	"""
	for mes in mesesR32:
		if file[3] == mes:
			return (mesesR32[mes].capitalize(),int(meses[mesesR32[mes]]),mes)
	else: raise ErrorMesBarra(file)
		
def getData():
	"""
	devuelve un diccionario con los dats, errs y tension consigna de cada barra
	es llamada por main()
	{barra:{'dats':[],'errs':[],'tensionConsigna':7900.......},}
	"""
	global semestre,año
	dictBarras = {}
	barras,dictBarras,dats,errs,barrasMesDatErr,resultados = {},{},{},{},{},{}
	
	# Detecta los meses que hay para analizar
	for root, folders, files in walk(dir):
		carpetas = folders
		break
	foldersMeses = [x for x in carpetas if x.lower() in meses.keys()]
	
	# En base a los meses que hay, detecta el semestre
	test = 0
	for mes in foldersMeses:
		if mes.lower() in ['enero','febrero','marzo','abril','mayo','junio']: test+=1
	if test == 6: semestre = '1'
	elif test == 1: semestre = '2'
	
	try:
		año = search('20\d\d',dir).group()
	except AttributeError:
		año = str(localtime().tm_year)
		if '-y' not in argv:
			print('Año del semestre:> {}\r'.format(año),end='Año del semestre:> ')
			año = input()
		else: año = argv[argv.index('-y')+1]
		if not año: año = str(localtime().tm_year)
	
	# Arma diccionario con mes y archivos en la carpeta de ese mes
	filesPorMes = dict(zip(foldersMeses,[getFiles(r'{}/{}/Archivos Enviados'.format(dir,mes))[0] for mes in foldersMeses]))
	
	# Arma diccionario con mes y excel de resultados para ese mes
	for mes in foldersMeses:
		for f in filesPorMes[mes]:
			if 'Resultados Mediciones' in f and f.endswith('.xlsx'):
				resultados[mes] = f
				break
		else:
			salir('No se encontro el informe de resultados en la carpeta de {}'.format(mes.capitalize()))
	if 'Diciembre' in resultados: resultados['diciembreplus'] = ''
	if 'Enero' in resultados: resultados['eneroplus'] = ''
	
	mesesArgLocal = set([list(mesesArg)[0]-1]+list(mesesArg)+[list(mesesArg)[-1]+1]).intersection(set([int(meses[x.lower()]) for x in resultados.keys()]))
	keysMeses = set(sorted([x.capitalize() for x in resultados.keys()],key=lambda x: int(meses[x.lower()]))).intersection(set((mesesRev[str(z).zfill(2)].capitalize() for z in mesesArgLocal)))
	# Arma una lista de meses, con cada barra y todos sus datos
	listaMeses = []
	for m in keysMeses:
		mes = Mes(m)
		resultadosMedicionesBarrasMes = openpyxl.load_workbook('{}/{}/Archivos Enviados/{}'.format(dir,m,resultados[m]))[sheets['BARRAS']]
		for fila in resultadosMedicionesBarrasMes:
			if fila[0].value and 'C/D' in fila[0].value:
				barra = Barra(m,fila)
				if barrasArg != [] and barra.id not in barrasArg: continue
				if search('(?i)fa(?:ll)?(?:ida)?',barra.resultadoInformado):
					print('Se saltea barra {} para el mes de {} por ser informada como fallida.'.format([barra.id,barra.puntoDeMedicion,barra.nombre],mes.mes))
					barra.preguntar = True
					continue
				if not barra.archivo:
					print('Se saltea barra {} para el mes de {} porque no hay archivo R32.'.format([barra.id,barra.puntoDeMedicion,barra.nombre],mes.mes))
					barra.preguntar = True
					continue
				barra.dat = Dat('{}/{}/Archivos Enviados/'.format(dir,m,),barra.archivo+'.dat',mes,barra)
				barra.err = Err('{}/{}/Archivos Enviados/'.format(dir,m,),barra.archivo+'.err',mes,barra)
				if barra.dat.problemas or barra.err.problemas: barra.preguntar = True
				mes.barras.append(barra)
		dictBarras = dict(zip(['_'.join((barra.id,str(barra.puntoDeMedicion))) for barra in mes.barras],mes.barras))
		mes.barras = dictBarras
		listaMeses.append(mes)
	if not all(all(not x.barras[y].preguntar for y in x.barras) for x in listaMeses):
		choice = input('Continuar?')
		if choice not in choices['yes']: salir()
	
	dictMeses = dict(zip([mes.mesInt for mes in listaMeses],listaMeses))
	return dictMeses

def cargarUnitarios():
	"""
	cargarUnitarios()
	Devuelve un diccionario con la siguiente estructura:
	{DESVIACION:VALOR_UNITARIO,DESVIACION:VALOR_UNITARIO,...}
	"""
	global valoresUnitariosPrevios,valoresUnitariosPost,fechaCambioUnitarios
	mesCambioUnitarios = '01' if semestre == '1' else '07'
	cambioStamp = '23/{}/{}'.format(mesCambioUnitarios,año[2:])
	listaUnitarios = list(reversed(unitariosDoc))
	for doc in listaUnitarios:
		try: stampStr = search('\d{1,2}-\d{1,2}-\d{4}',doc).group()
		except: continue
		stamp = strptime(stampStr,'%d-%m-%Y')
		fechaCambioUnitarios = strptime(cambioStamp,'%d/%m/%y')
		if mktime(fechaCambioUnitarios) == mktime(stamp):
			stampPrevio = strptime(search('\d{1,2}-\d{1,2}-\d{4}',listaUnitarios[listaUnitarios.index(doc)+1]).group(),'%d-%m-%Y')
			docFiles = [Document(doc),Document(listaUnitarios[listaUnitarios.index(doc)+1])] # esto puede dar problemas en casos de borde
			factores = [factoresDeInversion['ESJ'][sorted(list(filter(lambda x: mktime(stamp)+3600 > mktime(strptime(x,'%d/%m/%y')),factoresDeInversion['ESJ'])),key = lambda x: mktime(strptime(x,'%d/%m/%y')))[-1]],factoresDeInversion['ESJ'][sorted(list(filter(lambda x: mktime(stampPrevio)+3600 > mktime(strptime(x,'%d/%m/%y')),factoresDeInversion['ESJ'])),key = lambda x: mktime(strptime(x,'%d/%m/%y')))[-1]]]
			print("""
			Año: {}\t|\tSemestre: {}
			Usando unitarios vigentes hasta y desde el {}
			Usando factores de inversion = {}
			""".format(año,semestre,strftime('%d/%m/%Y',stamp),factores))
			break
	else: salir(mensaje='\nNo se encontraron los unitarios para {}'.format(strftime('%d/%m/%Y',fecha)))
	rangos = [.05,.06,.07,.08,.09,.1,.11,.12,.13,.14,.15,.16,.18]
	valoresUnitariosPost = dict(zip(rangos[2:],[float(x.text.replace(',','.'))*factores[0] for x in docFiles[0].tables[1].columns[3].cells[1:]]))
	valoresUnitariosPrevios = dict(zip(rangos[2:],[float(x.text.replace(',','.'))*factores[-1] for x in docFiles[-1].tables[1].columns[3].cells[1:]]))

def sincronizar(data):
	"""
	Esta funcion ordena los registros y los asocia al mes correspondiente.
	Ej: Los registros del dat de diciembre que son de enero van a parar a un dat de enero.
	"""
	
	listaMeses = sorted(data.keys())
	for id,mes in enumerate(listaMeses):							# MES
		listaBarras = sorted(data[mes].barras.keys())
		for barra in listaBarras:									# BARRA
			c = 0
			for vecino in list(seekExpansivo(listaMeses,id))[1:]: 	# MES
				if not data[vecino].barras.__contains__(barra): continue
				for registro in data[vecino].barras[barra].dat.intrusos:
					if registro.fechaTupla.tm_mon == mes and mktime(registro.fechaTupla) not in data[mes].barras[barra].dat.timestamps:
						data[mes].barras[barra].dat.registros.append(registro)
						data[mes].barras[barra].dat.timestamps.append(mktime(registro.fechaTupla))
						c+=1
				for corte in data[vecino].barras[barra].err.intrusos:
					if corte.fechaTupla.tm_mon == mes:
						data[mes].barras[barra].err.registros.append(corte)
						c+=1
			if c:
				data[mes].barras[barra].dat.registros.sort(key = lambda x: mktime(x.fechaTupla))
				data[mes].barras[barra].dat.setTimeStamps(data[mes].barras[barra],sincronizar=True)
	return data

def inRange(numeros,lInf,lSup):
	"""
	Se fija si cada valor de la lista "numeros" está dentro de los rangos que se le entregan a la funcion.
	Si todos los valores están dentro de todos los rangos, devuelve True
	Si algún valor está por fuera de algún rango, devuelve False
	"""
	for numero in numeros:
		if numero < lSup and numero > lInf: continue
		else: break
	else: return True
	return False

def armarExcel(mes,barra):
	"""
	Genera el excel resumen del procesamiento de cada medicion y lo guarda en el mismo directorio
	donde están los dats.
	"""
	
	plantilla = plantillaBarras
	celdasPlantillaBarras = {'A13': barra.direccion,
							'B10': barra.tensionNominal,
							'B11': barra.localidad,
							'B5': barra.mes,
							'B6': año,
							'B7': barra.idNumero,
							'B8': barra.nombre,
							'B9': barra.puntoDeMedicion,
							'C16': barra.dat.totalRegistros,
							'C17': barra.dat.totalRegistrosSobretension,
							'C18': barra.dat.totalRegistrosSubtension,
							'C19': barra.dat.totalRegistrosPenalizados,
							'C20': barra.dat.energiaTotal,
							'C21': barra.dat.energiaSobretension,
							'C22': barra.dat.energiaSubtension,
							'C23': barra.dat.energiaPenalizada,
							'C24': barra.dat.multaFueraDeRango,
							'C25': barra.dat.thdTotal,
							'C26': barra.dat.flicker,
							'C27': barra.dat.promedioTension,
							'C28': barra.dat.tensionMaxima,
							'C29': barra.dat.tensionMinima,
							'C3': 'Energía San Juan S.A.',
							'C30': barra.dat.apartamientoPromedio,
							'D13': barra.dat.resultado,
							'D16': barra.dat.totalRegistrosF1,
							'D17': barra.dat.totalRegistrosSobretensionF1,
							'D18': barra.dat.totalRegistrosSubtensionF1,
							'D19': barra.dat.totalRegistrosPenalizadosF1,
							'D20': barra.dat.energiaTotalF1,
							'D27': barra.dat.promedioTensionF1,
							'D28': barra.dat.tensionMaximaF1,
							'D29': barra.dat.tensionMinimaF1,
							'D30': barra.dat.apartamientoPromedioF1,
							'E11': barra.dat.tv,
							'E12': barra.dat.ti,
							'E16': barra.dat.totalRegistrosF2,
							'E17': barra.dat.totalRegistrosSobretensionF2,
							'E18': barra.dat.totalRegistrosSubtensionF2,
							'E19': barra.dat.totalRegistrosPenalizadosF2,
							'E20': barra.dat.energiaTotalF1,
							'E27': barra.dat.promedioTensionF2,
							'E28': barra.dat.tensionMaximaF2,
							'E29': barra.dat.tensionMinimaF2,
							'E30': barra.dat.apartamientoPromedioF2,
							'F10': barra.r32,
							'F16': barra.dat.totalRegistrosF3,
							'F17': barra.dat.totalRegistrosSobretensionF3,
							'F18': barra.dat.totalRegistrosSubtensionF3,
							'F19': barra.dat.totalRegistrosPenalizadosF3,
							'F20': barra.dat.energiaTotalF1,
							'F27': barra.dat.promedioTensionF3,
							'F28': barra.dat.tensionMaximaF3,
							'F29': barra.dat.tensionMinimaF3,
							'F30': barra.dat.apartamientoPromedioF3,
							'F6': barra.dat.fechaInicio,
							'F8': barra.dat.fechaFin,
							'F9': barra.dat.horaFin,
							'F7': barra.dat.horaInicio,}
	wb = openpyxl.load_workbook(plantilla)
	ws = wb['Plantilla']

	for celda in celdasPlantillaBarras:
		ws[celda].value = celdasPlantillaBarras[celda]
		
	for fase in range(1,4):
		tensiones = [[]]+[[getattr(reg,'V{}'.format(fase))] for reg in barra.dat.registros]
		armarGrafico(fase,tensiones,ws,barra)

	filename = ''.join(('cd',barra.idNumero,str(barra.puntoDeMedicion),semestre,año[-2:],'.xlsx'))
	filepath = '/'.join([dir.replace('\\','/'),mes,filename])
	try:
		wb.save(filepath)
	except(PermissionError):
		print('{} Está siendo usado por otro proceso, o está abierto. Cerralo y apretá ENTER para continuar...'.format(filename))
		input()
		wb.save(filepath)
	return filepath
	
def armarGrafico(fase,tensiones,ws,barra):
	"""
	Arma los gráficos de excel que van a ir en los archivos
	"""
	posicion = posicionesGraficosBarras[fase]
	c=1
	for V in tensiones[1:]:
		ws.cell(row=c, column=17+fase, value=V[0])
		c+=1
		
	###### Dibuja los limites superior e inferior
	c=1
	for _ in tensiones[1:]:
		ws.cell(row=c, column=22, value=barra.lSupPen)
		c+=1
	c=1
	for _ in tensiones[1:]:
		ws.cell(row=c, column=23, value=barra.lInfPen)
		c+=1
	######

	c1 = LineChart()
	c1.title = "Fase {}".format(fase)
	c1.style = 13
	c1.y_axis.title = 'Tensión [V]'
	c1.x_axis.title = 'Registros'

	vList = [x[0] for x in tensiones[1:]]
	minimo = min([min(vList),barra.lInfPen])
	maximo = max([max(vList),barra.lSupPen])
	c1.y_axis.scaling.min = minimo*0.95
	c1.y_axis.scaling.max = maximo*1.05
	
	font = Font(typeface='Verdana')
	size = 1250
	cp = CharacterProperties(latin=font, sz=size, b=False)
	c1.x_axis.title.tx.rich.p[0].pPr = ParagraphProperties(defRPr=cp)
	c1.y_axis.title.tx.rich.p[0].pPr = ParagraphProperties(defRPr=cp)
	c1.title.tx.rich.p[0].pPr = ParagraphProperties(defRPr=cp)
	c1.x_axis.txPr = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=cp), endParaRPr=cp)])
	c1.y_axis.txPr = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=cp), endParaRPr=cp)])
	c1.title.txPr = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=cp), endParaRPr=cp)])
	c1.x_axis.tickLblSkip = 200
	c1.x_axis.scaling.min = 0
	c1.legend = None

	data = Reference(ws, min_col=17+fase, min_row=1, max_col=17+fase, max_row=len(tensiones))
	limiteSuperior = Reference(ws, min_col=22, min_row=1, max_col=22, max_row=len(tensiones))
	limiteInferior = Reference(ws, min_col=23, min_row=1, max_col=23, max_row=len(tensiones))
	c1.add_data(data, titles_from_data=True)
	c1.add_data(limiteSuperior, titles_from_data=True)
	c1.add_data(limiteInferior, titles_from_data=True)

	s2 = c1.series[0]
	s2.graphicalProperties.line.width = Pt(1)
	s2.graphicalProperties.line.solidFill = "000000"
	c1.series[1].graphicalProperties.line.width = Pt(2.2)
	c1.series[1].graphicalProperties.line.solidFill = "000000"
	c1.series[2].graphicalProperties.line.width = Pt(2.2)
	c1.series[2].graphicalProperties.line.solidFill = "000000"

	c1.height = 12
	c1.width = 24
	ws.add_chart(c1, posicion)
	
def errSemestral(barra):
	global filaErrSemestral
	fila = filaErrSemestral
		
	while True:
		data = (yield)
		if data == 'close': break
		inicio,fin = data
		if inicio.duracion:
			if barra.err.serie1612:
				duracion = inicio.duracion
			else:
				horas, rem = divmod(int(inicio.duracion), 3600) 
				minutos, segundos = divmod(rem, 60)
				duracion = '{:0>2}:{:0>2}:{:05.2f}'.format(horas,minutos,segundos)
		else: duracion = None
		workSheetErrSemestral['A%d'%fila].value = inicio.fecha
		workSheetErrSemestral['B%d'%fila].value = inicio.hora
		workSheetErrSemestral['C%d'%fila].value = inicio.detalle
		workSheetErrSemestral['D%d'%fila].value = ' '.join((barra.id,barra.nombre))
		if inicio.serie1612:
			workSheetErrSemestral['E%d'%fila].value = ' '.join([inicio.fecha,inicio.hora]) if 'inicio' in inicio.detalle.lower() else ''
			workSheetErrSemestral['F%d'%fila].value = ' '.join([fin.fecha,fin.hora]) if 'inicio' in inicio.detalle.lower() else ''
		else:
			workSheetErrSemestral['E%d'%fila].value = ' '.join([inicio.fecha,inicio.hora]) if 'corte' in inicio.detalle.lower() else ''
			workSheetErrSemestral['F%d'%fila].value = ' '.join([fin.fecha,fin.hora]) if 'corte' in inicio.detalle.lower() else ''
		workSheetErrSemestral['G%d'%fila].value = duracion
		workSheetErrSemestral['H%d'%fila].value = inicio.errName
		workSheetErrSemestral['A%d'%(fila+1)].value = fin.fecha
		workSheetErrSemestral['B%d'%(fila+1)].value = fin.hora
		workSheetErrSemestral['C%d'%(fila+1)].value = fin.detalle
		workSheetErrSemestral['D%d'%(fila+1)].value = ' '.join((barra.id,barra.nombre))
		workSheetErrSemestral['H%d'%(fila+1)].value = fin.errName
		fila += 2
		filaErrSemestral = fila

def genBarrasSemestral():
	global filaDatSemestral,mesDatSemestral
	fila = filaDatSemestral
	fila = 12
		
	while True:
		data = (yield)
		if data == 'close': break
		elif type(data) == int:
			mesDatSemestral = data
			mesStr = mesesRev[str(mesDatSemestral).zfill(2)]
			fila+=3
			workSheetBarrasSemestral['A%d'%fila].value = ' '.join([mesStr.capitalize(),año])
			workSheetBarrasSemestral.row_dimensions[fila].fill = fillColor(1)
			workSheetBarrasSemestral['A%d'%fila].font = cellFontStyle('bold')
			workSheetBarrasSemestral['A%d'%fila].fill = fillColor(1)
			
			fila += 2
			filaDatSemestral = fila
			continue
		elif type(data) == Barra:
			barra = data
			registrosPorDia = 1440/float(barra.dat.periodo)
			if barra.puntoDeMedicion == 1:
				workSheetBarrasSemestral['A%d'%fila].value = barra.nombre
				workSheetBarrasSemestral['B%d'%fila].value = barra.id
			workSheetBarrasSemestral['C%d'%fila].value = str(barra.puntoDeMedicion)
			workSheetBarrasSemestral['D%d'%fila].value = str(monthrange(int(año),mesDatSemestral)[1])
			workSheetBarrasSemestral['F%d'%fila].value = round(barra.dat.totalRegistros/registrosPorDia,2)
			workSheetBarrasSemestral['G%d'%fila].value = round(barra.dat.totalRegistros/registrosPorDia,2)
			workSheetBarrasSemestral['H%d'%fila].value = round(barra.dat.totalRegistros/registrosPorDia,2)
			workSheetBarrasSemestral['I%d'%fila].value = round(barra.dat.totalRegistros/registrosPorDia,2)
			workSheetBarrasSemestral['K%d'%fila].value = barra.tensionNominal
			workSheetBarrasSemestral['M%d'%fila].value = barra.dat.tensionMaximaF1
			workSheetBarrasSemestral['N%d'%fila].value = barra.dat.tensionMaximaF2
			workSheetBarrasSemestral['O%d'%fila].value = barra.dat.tensionMaximaF3
			workSheetBarrasSemestral['Q%d'%fila].value = barra.dat.tensionMinimaF1
			workSheetBarrasSemestral['R%d'%fila].value = barra.dat.tensionMinimaF2
			workSheetBarrasSemestral['S%d'%fila].value = barra.dat.tensionMinimaF3
			workSheetBarrasSemestral['U%d'%fila].value = barra.dat.totalRegistrosPenalizadosF1
			workSheetBarrasSemestral['V%d'%fila].value = barra.dat.totalRegistrosPenalizadosF2
			workSheetBarrasSemestral['W%d'%fila].value = barra.dat.totalRegistrosPenalizadosF3
			workSheetBarrasSemestral['X%d'%fila].value = barra.dat.totalRegistrosPenalizados
			workSheetBarrasSemestral['Y%d'%fila].value = round(barra.dat.totalRegistrosPenalizados/barra.dat.totalRegistros*100,2)
			fila+=1
			filaDatSemestral=fila

def clean(tipoProceso):
	if tipoProceso != 2:
		wbDat.save(filenameD)
	if tipoProceso != 1:
		wbErr.save(filenameE)
	
def setup(tipoProceso):
	localDir = '\\'.join([dir,'Semestral'])
	if not path.isdir(localDir):
		mkdir(localDir)
		
	if tipoProceso != 2:
		global workSheetBarrasSemestral
		global filenameD
		global wbDat
		
		filenameD = '\\'.join([localDir,'Semestral.xlsx'])
		if path.isfile(filenameD):
			remove(filenameD)
		wbDat = openpyxl.load_workbook(plantillaSemestral)
		workSheetBarrasSemestral = wbDat['Resumen Barras']
		workSheetBarrasSemestral['A3'].value = '{} Semestre del {}'.format(semestre,año)
		
	if tipoProceso != 1:
		global workSheetErrSemestral
		global filenameE
		global wbErr
		
		filenameE = '\\'.join([localDir,'CortesSemestral.xlsx'])
		if path.isfile(filenameE):
			remove(filenameE)
		wbErr = openpyxl.load_workbook(plantillaCortesSemestral)
		workSheetErrSemestral = wbErr['Interrupciones']
		workSheetErrSemestral['D1'].value = '{} Semestre del {}'.format(semestre,año)

@manejarErrores
def main():
	global tipoProceso,mesesArg,barrasArg
	cortesSemestral = []
	listaFinal = []
	archivosExcel = []
	getDir()
	
	print('Procesar:\n1) Barras\n2) Cortes\n3) Barras y Cortes')
	tipoProceso = int(input(':> '))
	if tipoProceso not in range(1,4):
		print('Numero inválido')
		exit()
	print('Meses a procesar? (*,1,2,3,4,5,6,7,8,9,10,11,12)')
	mesesArg = {1,2,3,4,5,6,7,8,9,10,11,12}
	mesesIn = input(':> ').split(',')
	if mesesIn != ['*']: 
		mesesIn = set((int(x) for x in mesesIn))
		mesesArg = mesesArg.intersection(mesesIn)
	print('Barras a procesar? (*|[Numero de barra,...])')
	choice = input(':> ').strip().split(',')
	if choice == ['*']:	barrasArg = []
	else: barrasArg = ['C/D'+x.zfill(2) for x in choice]
	armarExcelSemestral = input('¿Armar un excel mensual para cada barra?\n(S/n)> ').lower()
	print()
	
	data = getData() # getData() detecta el semestre y el año en el que se está trabajando y los declara globalmente
	regsPorMes = dict(zip([int(x) for x in sorted(list(meses.values())) if int(x) in range(1,13)],list(map(lambda x: int(monthrange(int(año),int(x))[1]/15*1440),sorted(filter(lambda x : x in range(1,13),list(meses.values())))))))
	data = sincronizar(data)
	setup(tipoProceso)
	cargarUnitarios()
	
	listaMeses = sorted(data.keys())
	if semestre == '1' and 12 in listaMeses: listaMeses.remove(12)
	elif semestre == '2' and 6 in listaMeses: listaMeses.remove(6)
	listaMeses = set(listaMeses).intersection(mesesArg)
	
	if tipoProceso != 2:
		genDatSemestral = genBarrasSemestral()
		genDatSemestral.send(None)
	for mes in listaMeses:
		mes = data[mes]
		print(mes.mes)
		if tipoProceso != 2: genDatSemestral.send(mes.mesInt)
		listaBarras = sorted(mes.barras.keys(),key = lambda x: float(mes.barras[x].idNumero+'.'+str(mes.barras[x].puntoDeMedicion)))
		for barra in listaBarras:
			barra = mes.barras[barra]
			if tipoProceso != 2:
				print('\t',barra.id,barra.nombre,barra.dat.name)
				if barra.dat.problemas: continue
				barra.procesarDat()
				genDatSemestral.send(barra)
				if armarExcelSemestral in choices['yes']:
					armarExcel(mes.mes,barra)
	
			if tipoProceso != 1:
				print('\t',barra.id,barra.nombre,barra.err.name)
				if barra.err.problemas: continue
				barra.procesarErr()
				genErrSemestral = errSemestral(barra)
				genErrSemestral.send(None)
				for corte in barra.cortes:
					genErrSemestral.send(corte)
					
				try:genErrSemestral.send('close')
				except StopIteration:pass
	
	if tipoProceso != 2:
		try:genDatSemestral.send('close')
		except StopIteration:pass

	clean(tipoProceso)
	
if __name__ == '__main__':
	TIEMPO = time()
	main()
	elapsed = time()-TIEMPO
	mins,segs = divmod(elapsed,60)
	print(round(mins),' minutos',round(segs),'segundos')
	salir(mensaje='Listo...')