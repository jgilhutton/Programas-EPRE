import pyodbc
import subprocess
from docx import Document
from docx.shared import Pt
from re import sub,search
from time import mktime,strptime,strftime,localtime,time,gmtime
from rutas import dbCortes,dbUsuarios
from wordExe import wordExe
from Informacion import fechaFinRegistrosCortes

#VARIABLES
columnaORepo = False
años = 3
tablaUsuarios = 'suministros_sidac'

def getTablas():
	connCortes = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbCortes)
	cursorCortes = connCortes.cursor()
	tablas = [x[2] for x in cursorCortes.tables() if not x[2].startswith('MS')]
	tablas = sorted(tablas,key = lambda x: int(search('^\d{1,2}',x).group()))
	tablaCortes7, tablaCortes6, tablaCortes5, tablaCortes4, tablaCortes3, tablaCortes2, tablaCortes1 = tablas[::-1][:7]
	return tablaCortes7, tablaCortes6, tablaCortes5, tablaCortes4, tablaCortes3, tablaCortes2, tablaCortes1

def fechasCortes(años):
	"""
	Devuelve: [Fecha inicio, Fecha fin]
	"""
	segundosFinRegistrosCortes = mktime(strptime(fechaFinRegistrosCortes,'%d/%m/%y'))
	segundosInicioRegistrosCortes = segundosFinRegistrosCortes-(60*60*24*365*años) # - 6 semestres en segundos
	año = localtime(segundosInicioRegistrosCortes).tm_year
	mes = localtime(segundosInicioRegistrosCortes).tm_mon
	cappedTimeTuple = strptime(str(año),'%Y') if mes < 7 else strptime('%d %d'%(7,año),'%m %Y') # corre la fecha al inicio del semestre. si cae en el 1° semestre corre la fecha al primero de enero. 1 de julio para el 2° semestre
	print([strftime('%d/%m/%y',cappedTimeTuple),strftime('%d/%m/%y',localtime(segundosFinRegistrosCortes+86400))])
	return [strftime('%d/%m/%y',cappedTimeTuple),strftime('%d/%m/%y',localtime(segundosFinRegistrosCortes+86400))]

def procesarCortes(plantilla=None,tarifa=None,años=3):
	fechas = fechasCortes(años)

	if tarifa:
		connUsuarios = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbUsuarios)
		cursorUsuarios = connUsuarios.cursor()
		query = 'SELECT SETA,CLIENTE,TARIFA FROM [{}] WHERE CLIENTE = ?'.format(tablaUsuarios)
		cursorUsuarios.execute(query,(suministro,))
		resultado = cursorUsuarios.fetchall()
		tarifa = resultado[0][-1]
		seta = resultado[0][0]
	else: tarifa,seta = None,None
	connCortes = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=%s;' %dbCortes)
	cursorCortes = connCortes.cursor()
	#UNION SELECT [TABLA2].Suministro, [TABLA2].Interrupcion, [TABLA2].Inicio, [TABLA2].Final, [TABLA3].Id, [TABLA2].Motivo_EPRE FROM [TABLA2] WHERE (Suministro = ?) \
	query = 'SELECT [TABLA1].Suministro, [TABLA1].Interrupcion, [TABLA1].Orden_Reposicion, [TABLA1].Inicio, [TABLA1].Final, [TABLA1].Id, [TABLA1].Motivo_EPRE FROM [TABLA1] WHERE (Suministro = ? and Inicio >= DATEVALUE(?)) \
	UNION SELECT [TABLA2].Suministro, [TABLA2].Interrupcion, [TABLA2].Orden_Reposicion, [TABLA2].Inicio, [TABLA2].Final, [TABLA2].Id, [TABLA2].Motivo_EPRE FROM [TABLA2] WHERE (Suministro = ? and Inicio >= DATEVALUE(?)) \
	UNION SELECT [TABLA3].Suministro, [TABLA3].Interrupcion, [TABLA3].Orden_Reposicion, [TABLA3].Inicio, [TABLA3].Final, [TABLA3].Id, [TABLA3].Motivo_EPRE FROM [TABLA3] WHERE (Suministro = ? and Inicio >= DATEVALUE(?)) \
	UNION SELECT [TABLA4].Suministro, [TABLA4].Interrupcion, [TABLA4].Orden_Reposicion, [TABLA4].Inicio, [TABLA4].Final, [TABLA4].Id, [TABLA4].Motivo_EPRE FROM [TABLA4] WHERE (Suministro = ? and Inicio >= DATEVALUE(?)) \
	UNION SELECT [TABLA5].Suministro, [TABLA5].Interrupcion, [TABLA5].Orden_Reposicion, [TABLA5].Inicio, [TABLA5].Final, [TABLA5].Id, [TABLA5].Motivo_EPRE FROM [TABLA5] WHERE (Suministro = ? and Inicio >= DATEVALUE(?)) \
	UNION SELECT [TABLA6].Suministro, [TABLA6].Interrupcion, [TABLA6].Orden_Reposicion, [TABLA6].Inicio, [TABLA6].Final, [TABLA6].Id, [TABLA6].Motivo_EPRE FROM [TABLA6] WHERE (Suministro = ? and Inicio >= DATEVALUE(?)) \
	UNION SELECT [TABLA7].Suministro, [TABLA7].Interrupcion, [TABLA7].Orden_Reposicion, [TABLA7].Inicio, [TABLA7].Final, [TABLA7].Id, [TABLA7].Motivo_EPRE FROM [TABLA7] WHERE (Suministro = ? and Inicio <= DATEVALUE(?));'
	query = query.replace('TABLA1',tablaCortes1)
	query = query.replace('TABLA2',tablaCortes2)
	query = query.replace('TABLA3',tablaCortes3)
	query = query.replace('TABLA4',tablaCortes4)
	query = query.replace('TABLA5',tablaCortes5)
	query = query.replace('TABLA6',tablaCortes6)
	query = query.replace('TABLA7',tablaCortes7)
	cursorCortes.execute(query,(suministro,fechas[0],suministro,fechas[0],suministro,fechas[0],suministro,fechas[0],suministro,fechas[0],suministro,fechas[0],suministro,fechas[1]))
	cortes = cursorCortes.fetchall()
	cortes = sorted(cortes, key = lambda x: (x[3] - x[3].utcfromtimestamp(0)).total_seconds())
	
	if not cortes:
		print('No hay cortes para {}'.format(suministro))
	elif __name__ != '__main__':
		return fechas,cortes,tarifa,seta
	else:
		fila = 1
		for corte in cortes:
			# SUMINISTRO, INTERRUPCION, INICIO, FINAL, ID, MOTIVO EPRE
			plantilla.tables[0].rows[fila].cells[0].text = corte[0]
			plantilla.tables[0].rows[fila].cells[1].text = corte[1]
			plantilla.tables[0].rows[fila].cells[2].text = corte[3].strftime('%d/%m/%Y %H:%M')
			plantilla.tables[0].rows[fila].cells[3].text = corte[4].strftime('%d/%m/%Y %H:%M')
			duracion = round((corte[4]-corte[3]).total_seconds()/60)
			plantilla.tables[0].rows[fila].cells[4].text = str(duracion)
			plantilla.tables[0].rows[fila].cells[5].text = corte[6]
			c=0
			for _ in plantilla.tables[0].rows[fila].cells:
				for run in plantilla.tables[0].rows[fila].cells[c].paragraphs[0].runs:
					font = run.font
					font.size = Pt(11)
				plantilla.tables[0].rows[fila].cells[c].paragraphs[0].alignment = 1
				c+=1
			plantilla.tables[0].add_row()
			fila+=1
		plantilla.tables[0]._tbl.remove(plantilla.tables[0].rows[fila]._tr)
		plantilla.tables[0].style = 'Table Grid'
	return fechas
	t = tabla._element
	t.getparent().remove(t)
	t._p = t._element = None

def forImport(sum,tarifa,años=3):
	global suministro
	global tablaCortes7, tablaCortes6, tablaCortes5, tablaCortes4, tablaCortes3, tablaCortes2, tablaCortes1 
	tablaCortes7, tablaCortes6, tablaCortes5, tablaCortes4, tablaCortes3, tablaCortes2, tablaCortes1 = getTablas()
	suministro = sum
	return procesarCortes(tarifa=tarifa,años=años)
	
if __name__ == '__main__':
	tablaCortes7, tablaCortes6, tablaCortes5, tablaCortes4, tablaCortes3, tablaCortes2, tablaCortes1 = getTablas()

	print('Suministros de reclamo:\n')
	print('XXXXXXXXXXX\r',end='')
	suministro = input()
	print('Esto puede tardar un rato...')
	plantilla = Document('Recursos/Cortes 3 años.docx')
	fechas = procesarCortes(plantilla=plantilla,años=años)

	plantilla.save('Recursos/Cprocesado 3 años.docx')

	subprocess.Popen('"{}" "Recursos/Cprocesado 3 años.docx"'.format(wordExe),shell = True)

	print('\nListo...')
